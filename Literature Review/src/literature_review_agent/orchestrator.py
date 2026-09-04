"""The main orchestrator: runs the eleven pipeline stages, resumably.

Each stage reads its input from disk and writes its output back, then records a
checkpoint. Re-running the pipeline after any interruption — a dropped
connection, an API limit, a killed process — skips completed stages and resumes
long stages (download, extract, analyse) at the item they stopped on.

Every artefact produced is published through :class:`StorageManager`, so it is
uploaded to Google Drive and verified before being counted as delivered. A
failed upload keeps the local copy and stays retryable; it never aborts the run.
"""

from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

from .analysis import analyse_records
from .citation_manager import CitationManager, audit_citations
from .config import Settings, load_settings, missing_keys
from .deduplicator import deduplicate
from .downloader import download_records, failure_rows
from .evidence_ledger import EvidenceLedger
from .excel_report import (
    MatrixInputs,
    TopicColumn,
    build_citation_audit_workbook,
    build_evidence_workbook,
    build_workbook,
)
from .job_manager import Job
from .keyword_generator import build_keyword_strategy, write_keyword_outputs
from .llm import llm_available, unavailable_reason
from .logging_setup import get_logger, setup_logging
from .metadata import enrich_records, score_relevance
from .pdf_extractor import extract_records
from .q1_verifier import load_ranking_table, select_records, verify_records, write_pending_q1_csv
from .schemas import (
    STAGE_ORDER,
    CitationAuditRow,
    DownloadStatus,
    KeywordStrategy,
    PaperAnalysis,
    PaperRecord,
    Q1Mode,
    Q1Status,
    StageName,
    StageStatus,
)
from .search import filter_records, run_search
from .storage import StorageManager
from .synthesis import build_synthesis
from .utils import append_jsonl, read_json, write_json
from .verification import run_verification, write_unresolved_issues_csv
from .word_reports import (
    ReportContext,
    build_all_reports,
    build_unable_to_download,
    build_verification_report,
)

LOG = get_logger("orchestrator")


class StageError(RuntimeError):
    """Raised when a stage cannot proceed because a prerequisite is missing."""


@dataclass
class StageOutcome:
    """What one stage produced."""

    stage: StageName
    status: StageStatus
    message: str = ""
    artefacts: list[Path] = field(default_factory=list)
    counters: dict[str, int] = field(default_factory=dict)


@dataclass
class PipelineState:
    """The working set shared between stages, loaded from and saved to disk."""

    strategy: KeywordStrategy | None = None
    raw_records: list[PaperRecord] = field(default_factory=list)
    records: list[PaperRecord] = field(default_factory=list)
    analyses: dict[str, PaperAnalysis] = field(default_factory=dict)
    ledger: EvidenceLedger = field(default_factory=EvidenceLedger)
    audit_rows: list[CitationAuditRow] = field(default_factory=list)
    merge_events: list[dict[str, Any]] = field(default_factory=list)
    discovered_count: int = 0


class Orchestrator:
    """Coordinates the whole workflow for one job."""

    def __init__(
        self,
        job: Job,
        *,
        settings: Settings | None = None,
        storage: StorageManager | None = None,
        enable_drive: bool | None = None,
    ) -> None:
        self.job = job
        self.settings = settings or job.settings
        setup_logging(
            job.paths.pipeline_log,
            json_file=job.paths.logs / "pipeline.jsonl",
        )
        self.storage = storage or StorageManager(
            self.settings,
            job.paths,
            job_date=job.config.job_date,
            topic_slug=job.config.topic_slug,
            enable_drive=enable_drive,
        )
        self.state = PipelineState()

    # -- state persistence ---------------------------------------------

    def load_state(self) -> PipelineState:
        """Rehydrate the working set from the job's saved intermediate files."""
        paths = self.job.paths
        state = self.state

        if raw := read_json(paths.keyword_strategy_file):
            try:
                state.strategy = KeywordStrategy.model_validate(raw)
            except Exception as exc:  # noqa: BLE001
                LOG.warning(f"Saved keyword strategy unreadable ({exc}); it will be rebuilt.")

        for path, target in (
            (paths.raw_search_results_file, "raw_records"),
            (paths.records_file, "records"),
        ):
            rows = read_json(path, []) or []
            parsed: list[PaperRecord] = []
            for row in rows:
                try:
                    parsed.append(PaperRecord.model_validate(row))
                except Exception as exc:  # noqa: BLE001 - skip one bad row, keep the rest
                    LOG.debug(f"Skipping unreadable record in {path.name}: {exc}")
            setattr(state, target, parsed)

        for row in read_json(paths.analyses_file, []) or []:
            try:
                analysis = PaperAnalysis.model_validate(row)
                state.analyses[analysis.record_id] = analysis
            except Exception as exc:  # noqa: BLE001
                LOG.debug(f"Skipping unreadable analysis: {exc}")

        if paths.evidence_file.exists():
            state.ledger = EvidenceLedger.load(paths.evidence_file)

        state.merge_events = read_json(paths.logs / "merge_audit.json", []) or []
        counters = self.job.checkpoints.stage(StageName.SEARCH).counters
        state.discovered_count = counters.get("raw_records", len(state.raw_records))
        return state

    def save_records(self) -> Path:
        """Persist the canonical record set."""
        return write_json(
            self.job.paths.records_file,
            [r.model_dump(mode="json") for r in self.state.records],
        )

    def save_analyses(self) -> Path:
        """Persist the analysis records."""
        return write_json(
            self.job.paths.analyses_file,
            [a.model_dump(mode="json") for a in self.state.analyses.values()],
        )

    # -- publishing -----------------------------------------------------

    def publish(self, paths: list[Path], destination: str) -> list[Path]:
        """Upload artefacts to Drive and return the ones that were published."""
        published: list[Path] = []
        for path in paths:
            if path is None:
                continue
            outcome = self.storage.publish(Path(path), destination)
            published.append(Path(path))
            if outcome.drive_enabled and not outcome.verified:
                LOG.warning(
                    f"{Path(path).name} is staged locally but not yet verified in "
                    f"Drive ({outcome.status}). It stays queued for drive-sync."
                )
        return published

    # -- stage runner ---------------------------------------------------

    def run_stage(self, stage: StageName, *, force: bool = False) -> StageOutcome:
        """Run one stage, honouring its checkpoint unless *force* is set."""
        if not force and self.job.is_stage_complete(stage):
            LOG.info(f"Stage '{stage.value}' already complete; skipping.")
            return StageOutcome(stage, StageStatus.COMPLETE, "already complete")

        runner: Callable[[], StageOutcome] = getattr(self, f"_stage_{stage.value}")
        self.job.start_stage(stage)
        try:
            outcome = runner()
        except StageError as exc:
            self.job.fail_stage(stage, str(exc))
            LOG.error(f"Stage '{stage.value}' cannot proceed: {exc}")
            return StageOutcome(stage, StageStatus.FAILED, str(exc))
        except Exception as exc:  # noqa: BLE001 - record and keep the checkpoint
            self.job.fail_stage(stage, f"{type(exc).__name__}: {exc}")
            LOG.exception(f"Stage '{stage.value}' failed; progress has been checkpointed.")
            return StageOutcome(stage, StageStatus.FAILED, f"{type(exc).__name__}: {exc}")

        if outcome.status == StageStatus.SKIPPED:
            self.job.skip_stage(stage, outcome.message)
        else:
            self.job.complete_stage(
                stage,
                message=outcome.message,
                artefacts=outcome.artefacts,
                counters=outcome.counters,
            )
        return outcome

    def run(
        self,
        *,
        stages: list[StageName] | None = None,
        force: bool = False,
        stop_on_failure: bool = True,
    ) -> list[StageOutcome]:
        """Run the pipeline (or a subset of stages) from where it left off."""
        self.load_state()
        self.storage.ensure_remote_tree()

        outcomes: list[StageOutcome] = []
        for stage in stages or list(STAGE_ORDER):
            outcome = self.run_stage(stage, force=force)
            outcomes.append(outcome)
            if outcome.status == StageStatus.FAILED and stop_on_failure:
                LOG.error(
                    f"Stopping at '{stage.value}'. Fix the cause and re-run "
                    f"'resume' — completed stages will be skipped."
                )
                break
        return outcomes

    # ------------------------------------------------------------------
    # Stage 1: keywords
    # ------------------------------------------------------------------

    def _stage_keywords(self) -> StageOutcome:
        """Build and write the keyword strategy."""
        strategy = build_keyword_strategy(self.job.config, self.settings)
        self.state.strategy = strategy
        write_json(self.job.paths.keyword_strategy_file, strategy.model_dump(mode="json"))
        artefacts = write_keyword_outputs(strategy, self.job.config, self.job.paths.keywords)
        self.publish(artefacts, "keywords")
        self.publish([self.job.paths.keyword_strategy_file], "logs")

        return StageOutcome(
            StageName.KEYWORDS,
            StageStatus.COMPLETE,
            message=(
                f"{len(strategy.terms)} terms and {len(strategy.search_strings)} "
                f"search strings ({strategy.generator})."
            ),
            artefacts=artefacts,
            counters={
                "terms": len(strategy.terms),
                "search_strings": len(strategy.search_strings),
                "concepts": len(strategy.main_concepts),
            },
        )

    # ------------------------------------------------------------------
    # Stage 2: search
    # ------------------------------------------------------------------

    def _stage_search(self) -> StageOutcome:
        """Query every available source and save the raw records."""
        if self.state.strategy is None:
            raise StageError(
                "No keyword strategy is available. Run the keywords stage first."
            )

        outcome = run_search(self.state.strategy, self.job.config, self.settings)
        for log in outcome.logs:
            append_jsonl(self.job.paths.search_log, log.to_dict())

        kept, filter_counters = filter_records(
            outcome.records, self.job.config, self.settings
        )
        self.state.raw_records = kept
        self.state.discovered_count = len(outcome.records)
        write_json(
            self.job.paths.raw_search_results_file,
            [r.model_dump(mode="json") for r in kept],
        )
        self.publish(
            [self.job.paths.raw_search_results_file, self.job.paths.search_log], "logs"
        )

        if not kept:
            message = (
                "No records were returned by any source. Check the network connection, "
                "widen the year range, or broaden the keywords."
            )
            LOG.warning(message)

        return StageOutcome(
            StageName.SEARCH,
            StageStatus.COMPLETE,
            message=(
                f"{len(outcome.records)} raw records from "
                f"{len(outcome.sources_used)} source(s); {len(kept)} passed screening."
            ),
            artefacts=[self.job.paths.raw_search_results_file],
            counters={
                "raw_records": len(outcome.records),
                "after_screening": len(kept),
                "sources_used": len(outcome.sources_used),
                "sources_skipped": len(outcome.sources_skipped),
                "queries_run": len(outcome.logs),
                **filter_counters,
            },
        )

    # ------------------------------------------------------------------
    # Stage 3: deduplicate
    # ------------------------------------------------------------------

    def _stage_deduplicate(self) -> StageOutcome:
        """Merge duplicates, enrich metadata, and score relevance."""
        source = self.state.raw_records or self.state.records
        if not source:
            return StageOutcome(
                StageName.DEDUPLICATE,
                StageStatus.SKIPPED,
                message="No records to deduplicate.",
            )

        result = deduplicate(source, self.settings)
        self.state.records = result.records
        self.state.merge_events = [event.to_dict() for event in result.merges]
        write_json(self.job.paths.logs / "merge_audit.json", self.state.merge_events)

        enrich_records(self.state.records, self.settings, limit=None)

        strategy = self.state.strategy
        concepts = strategy.main_concepts if strategy else []
        keywords = strategy.query_terms() if strategy else []
        for record in self.state.records:
            score_relevance(
                record,
                concepts,
                keywords,
                self.settings,
                current_year=self.job.config.year_to,
            )

        self.save_records()
        self.publish(
            [self.job.paths.records_file, self.job.paths.logs / "merge_audit.json"], "logs"
        )

        return StageOutcome(
            StageName.DEDUPLICATE,
            StageStatus.COMPLETE,
            message=(
                f"{len(source)} records reduced to {len(result.records)} unique "
                f"({result.total_merged} merged)."
            ),
            artefacts=[self.job.paths.records_file],
            counters=result.counters,
        )

    # ------------------------------------------------------------------
    # Stage 4: Q1 verification
    # ------------------------------------------------------------------

    def _stage_q1_verify(self) -> StageOutcome:
        """Attach quartile evidence to every record."""
        if not self.state.records:
            return StageOutcome(
                StageName.Q1_VERIFY, StageStatus.SKIPPED, message="No records to verify."
            )

        table = load_ranking_table(self.settings, self.job.config.ranking_file)
        counters = verify_records(self.state.records, self.settings, table=table)
        self.save_records()
        self.publish([self.job.paths.records_file], "logs")

        message = (
            f"{counters.get(Q1Status.VERIFIED_Q1.value, 0)} verified Q1, "
            f"{counters.get(Q1Status.UNVERIFIED.value, 0)} unverified."
        )
        if table is None:
            message += (
                " No ranking file was configured, so no quartile was assigned rather "
                "than guessed."
            )

        return StageOutcome(
            StageName.Q1_VERIFY,
            StageStatus.COMPLETE,
            message=message,
            artefacts=[self.job.paths.records_file],
            counters={k: v for k, v in counters.items() if v},
        )

    # ------------------------------------------------------------------
    # Stage 5: selection
    # ------------------------------------------------------------------

    def _stage_select(self) -> StageOutcome:
        """Apply the inclusion rules and the Q1 mode."""
        if not self.state.records:
            return StageOutcome(
                StageName.SELECT, StageStatus.SKIPPED, message="No records to select from."
            )

        result = select_records(self.state.records, self.job.config, self.settings)
        artefacts: list[Path] = []

        if result.pending_q1:
            pending_name = self.settings.selection.get(
                "pending_list_filename", "pending_q1_verification.csv"
            )
            pending_path = write_pending_q1_csv(
                result.pending_q1, self.job.paths.verification / pending_name, self.settings
            )
            artefacts.append(pending_path)
            self.publish([pending_path], "verification")

        self.save_records()
        self.publish([self.job.paths.records_file], "logs")

        message = f"{len(result.selected)} papers included."
        if result.pending_q1:
            message += (
                f" {len(result.pending_q1)} await manual quartile verification and were "
                "not treated as Q1."
            )
        if self.job.config.q1_mode == Q1Mode.ONLY and not result.selected:
            message += (
                " Q1-only mode selected nothing: no paper has verified Q1 evidence. "
                "Supply a ranking file, or use q1_mode=preferred."
            )

        return StageOutcome(
            StageName.SELECT,
            StageStatus.COMPLETE,
            message=message,
            artefacts=artefacts,
            counters=result.counters,
        )

    # ------------------------------------------------------------------
    # Stage 6: download
    # ------------------------------------------------------------------

    def _stage_download(self) -> StageOutcome:
        """Download every selected paper's legally accessible PDF."""
        selected = [r for r in self.state.records if r.selected]
        if not selected:
            return StageOutcome(
                StageName.DOWNLOAD,
                StageStatus.SKIPPED,
                message="No papers were selected for download.",
            )

        done = self.job.done_items(StageName.DOWNLOAD)

        def on_complete(record_id: str) -> None:
            """Checkpoint after each paper so a crash resumes mid-list."""
            self.job.mark_item_done(StageName.DOWNLOAD, record_id)
            self.save_records()

        report = download_records(
            self.state.records,
            self.job.paths.downloaded_papers,
            self.settings,
            self.job.config,
            partial_dir=self.job.paths.partial,
            log_path=self.job.paths.download_log,
            already_done=done,
            on_complete=on_complete,
        )

        # Publish each validated PDF to Drive.
        pdf_paths = [
            Path(r.local_path)
            for r in self.state.records
            if r.local_path and Path(r.local_path).exists()
        ]
        self.publish(pdf_paths, "downloaded_papers")

        # Manual-retrieval register.
        rows = failure_rows(self.state.records)
        register = build_unable_to_download(
            rows,
            self.job.config,
            self.job.paths.unable_to_download / "Unable_to_Download.docx",
        )
        self.publish([register], "unable_to_download")

        artefacts = self._write_manifests()
        self.publish(artefacts, "papers")
        self.save_records()
        self.publish([self.job.paths.records_file, self.job.paths.download_log], "logs")

        return StageOutcome(
            StageName.DOWNLOAD,
            StageStatus.COMPLETE,
            message=(
                f"{report.counters['downloaded']} downloaded, "
                f"{report.counters['failed']} failed, "
                f"{report.counters['skipped_no_legal_url']} had no authorised "
                "open-access URL."
            ),
            artefacts=[register, *artefacts],
            counters=report.counters,
        )

    def _write_manifests(self) -> list[Path]:
        """Write the CSV and JSON paper manifests."""
        paths = self.job.paths
        fields = [
            "record_id", "title", "authors", "year", "journal", "volume", "issue",
            "pages", "doi", "issn", "publisher", "document_type", "language",
            "citation_count", "open_access_status", "licence", "landing_page_url",
            "pdf_url", "discovery_source", "metadata_sources", "relevance_score",
            "selected", "q1_status", "quartile", "ranking_source", "download_status",
            "local_filename", "file_sha256", "file_bytes", "requires_ocr",
            "failure_reason", "verification_confidence", "notes",
        ]
        with paths.paper_manifest_csv.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=fields)
            writer.writeheader()
            for record in self.state.records:
                writer.writerow(
                    {
                        "record_id": record.record_id,
                        "title": record.title,
                        "authors": "; ".join(record.authors),
                        "year": record.year or "",
                        "journal": record.journal,
                        "volume": record.volume or "",
                        "issue": record.issue or "",
                        "pages": record.pages or record.article_number or "",
                        "doi": record.doi or "",
                        "issn": record.issn or "",
                        "publisher": record.publisher or "",
                        "document_type": record.document_type or "",
                        "language": record.language or "",
                        "citation_count": record.citation_count
                        if record.citation_count is not None
                        else "",
                        "open_access_status": record.open_access_status or "",
                        "licence": record.licence or "",
                        "landing_page_url": record.landing_page_url or "",
                        "pdf_url": record.pdf_url or "",
                        "discovery_source": record.discovery_source,
                        "metadata_sources": "; ".join(record.metadata_sources),
                        "relevance_score": round(record.relevance_score, 4),
                        "selected": "Yes" if record.selected else "No",
                        "q1_status": record.q1.verification_status.value,
                        "quartile": record.q1.quartile or "",
                        "ranking_source": record.q1.ranking_source or "",
                        "download_status": record.download_status.value,
                        "local_filename": record.local_filename or "",
                        "file_sha256": record.file_sha256 or "",
                        "file_bytes": record.file_bytes or "",
                        "requires_ocr": "Yes" if record.requires_ocr else "No",
                        "failure_reason": record.failure_reason or "",
                        "verification_confidence": record.verification_confidence,
                        "notes": record.notes,
                    }
                )
        write_json(
            paths.paper_manifest_json,
            [r.model_dump(mode="json") for r in self.state.records],
        )
        return [paths.paper_manifest_csv, paths.paper_manifest_json]

    # ------------------------------------------------------------------
    # Stage 7: extract
    # ------------------------------------------------------------------

    def _stage_extract(self) -> StageOutcome:
        """Extract page-marked text from every downloaded PDF."""
        downloaded = [
            r for r in self.state.records
            if r.download_status in (DownloadStatus.DOWNLOADED, DownloadStatus.ALREADY_PRESENT)
        ]
        if not downloaded:
            return StageOutcome(
                StageName.EXTRACT,
                StageStatus.SKIPPED,
                message="No PDFs are available to extract.",
            )

        def on_complete(record_id: str) -> None:
            """Checkpoint after each extraction."""
            self.job.mark_item_done(StageName.EXTRACT, record_id)
            self.save_records()

        results = extract_records(
            self.state.records,
            self.settings,
            self.job.paths.extracted_text,
            already_done=self.job.done_items(StageName.EXTRACT),
            on_complete=on_complete,
        )

        text_paths = [
            Path(r.extracted_text_path)
            for r in self.state.records
            if r.extracted_text_path and Path(r.extracted_text_path).exists()
        ]
        self.publish(text_paths, "extracted_text")
        self.save_records()
        self.publish([self.job.paths.records_file], "logs")

        readable = sum(1 for r in results.values() if r.success)
        needs_ocr = sum(1 for r in self.state.records if r.requires_ocr)
        return StageOutcome(
            StageName.EXTRACT,
            StageStatus.COMPLETE,
            message=(
                f"{readable} of {len(text_paths)} PDFs produced readable text; "
                f"{needs_ocr} need OCR."
            ),
            artefacts=text_paths[:20],
            counters={
                "extracted": len(text_paths),
                "readable": readable,
                "requires_ocr": needs_ocr,
            },
        )

    # ------------------------------------------------------------------
    # Stage 8: analyse
    # ------------------------------------------------------------------

    def _stage_analyse(self) -> StageOutcome:
        """Analyse every paper with extracted text."""
        candidates = [
            r for r in self.state.records
            if r.extracted_text_path and Path(r.extracted_text_path).exists()
        ]
        if not candidates:
            return StageOutcome(
                StageName.ANALYSE,
                StageStatus.SKIPPED,
                message="No extracted text is available to analyse.",
            )

        manager = CitationManager(
            style=self.job.config.citation_style,
            records=[r for r in self.state.records if r.selected] or list(self.state.records),
        )

        def on_complete(record_id: str) -> None:
            """Checkpoint after each analysis."""
            self.job.mark_item_done(StageName.ANALYSE, record_id)

        new_analyses = analyse_records(
            self.state.records,
            self.job.config,
            self.settings,
            citations={
                r.record_id: manager.reference(r.record_id) for r in self.state.records
            },
            already_done=self.job.done_items(StageName.ANALYSE),
            on_complete=on_complete,
        )
        self.state.analyses.update(new_analyses)
        self.save_analyses()
        self.publish([self.job.paths.analyses_file], "logs")

        high = sum(
            1 for a in self.state.analyses.values() if a.overall_confidence == "High"
        )
        return StageOutcome(
            StageName.ANALYSE,
            StageStatus.COMPLETE,
            message=(
                f"{len(self.state.analyses)} papers analysed "
                f"({high} with high confidence)."
            ),
            artefacts=[self.job.paths.analyses_file],
            counters={
                "analysed": len(self.state.analyses),
                "high_confidence": high,
                "llm_assisted": sum(
                    1 for a in self.state.analyses.values() if "Claude" in a.analyser
                ),
            },
        )

    # ------------------------------------------------------------------
    # Stage 9: evidence and citation verification
    # ------------------------------------------------------------------

    def _stage_verify_evidence(self) -> StageOutcome:
        """Build the evidence ledger and audit citations before reporting.

        Synthesis runs here so the ledger exists before the reports are written,
        which is what lets the citation audit gate the report stage.
        """
        if not self.state.analyses:
            return StageOutcome(
                StageName.VERIFY_EVIDENCE,
                StageStatus.SKIPPED,
                message="No analyses are available, so there is nothing to verify.",
            )

        included = [r for r in self.state.records if r.selected] or list(self.state.records)
        manager = CitationManager(style=self.job.config.citation_style, records=included)
        self.state.ledger = EvidenceLedger()
        synthesis = build_synthesis(
            included, self.state.analyses, manager, self.state.ledger, self.job.config
        )
        write_json(
            self.job.paths.synthesis_file,
            {
                "gaps": [g.model_dump(mode="json") for g in synthesis.gaps],
                "models": [m.model_dump(mode="json") for m in synthesis.models],
                "landscape": synthesis.landscape.model_dump(mode="json"),
                "contradictions": synthesis.contradictions,
            },
        )
        self.state.ledger.save(self.job.paths.evidence_file)

        self.state.audit_rows = audit_citations(manager, self.state.ledger.records)
        write_json(
            self.job.paths.logs / "citation_audit.json",
            [r.model_dump(mode="json") for r in self.state.audit_rows],
        )
        manager.write_exports(
            self.job.paths.references_bib, self.job.paths.references_ris
        )
        self.publish(
            [self.job.paths.references_bib, self.job.paths.references_ris], "papers"
        )
        self.publish(
            [
                self.job.paths.evidence_file,
                self.job.paths.synthesis_file,
                self.job.paths.logs / "citation_audit.json",
            ],
            "logs",
        )

        failures = sum(1 for r in self.state.audit_rows if r.outcome.value == "Fail")
        return StageOutcome(
            StageName.VERIFY_EVIDENCE,
            StageStatus.COMPLETE,
            message=(
                f"{len(self.state.ledger.records)} evidence records, "
                f"{len(synthesis.gaps)} gaps, {len(synthesis.models)} models; "
                f"{failures} citation failure(s)."
            ),
            artefacts=[self.job.paths.evidence_file, self.job.paths.synthesis_file],
            counters={
                "evidence_records": len(self.state.ledger.records),
                "gaps": len(synthesis.gaps),
                "models": len(synthesis.models),
                "citation_failures": failures,
            },
        )

    # ------------------------------------------------------------------
    # Stage 10: reports
    # ------------------------------------------------------------------

    def _stage_report(self) -> StageOutcome:
        """Write the Excel matrix and the five Word synthesis documents."""
        if not self.state.records:
            return StageOutcome(
                StageName.REPORT, StageStatus.SKIPPED, message="No records to report on."
            )

        included = [r for r in self.state.records if r.selected] or list(self.state.records)
        manager = CitationManager(style=self.job.config.citation_style, records=included)

        # Rebuild the ledger so report writers register their own claims.
        self.state.ledger = EvidenceLedger()
        synthesis = build_synthesis(
            included, self.state.analyses, manager, self.state.ledger, self.job.config
        )

        context = ReportContext(
            config=self.job.config,
            settings=self.settings,
            records=included,
            analyses=self.state.analyses,
            manager=manager,
            ledger=self.state.ledger,
            synthesis=synthesis,
        )
        word_paths = build_all_reports(context, self.job.paths.reports)
        self.publish(word_paths, "reports")

        # Re-audit with the documents' actual text, then persist the ledger.
        document_texts = _read_document_texts(word_paths)
        self.state.audit_rows = audit_citations(
            manager, self.state.ledger.records, document_texts=document_texts
        )
        self.state.ledger.save(self.job.paths.evidence_file)

        excel_path = build_workbook(
            MatrixInputs(
                config=self.job.config,
                records=included,
                analyses=self.state.analyses,
                gaps=synthesis.gaps,
                models=synthesis.models,
                landscape=synthesis.landscape,
                download_failures=failure_rows(self.state.records),
                search_log=list(_read_jsonl(self.job.paths.search_log)),
                citation_audit=self.state.audit_rows,
                citations=manager.citation_map(),
                topic_columns=self._topic_columns(),
            ),
            self.settings,
            self.job.paths.reports
            / self.settings.reporting.get("excel_filename", "Literature_Review_Matrix.xlsx"),
        )
        self.publish([excel_path], "reports")

        evidence_path = build_evidence_workbook(
            self.state.ledger.rows(), self.job.paths.verification / "Evidence_Ledger.xlsx"
        )
        audit_path = build_citation_audit_workbook(
            [
                {
                    "serial": index,
                    "in_text_citation": row.in_text_citation,
                    "record_id": row.record_id,
                    "doi": row.doi or "",
                    "title": row.title,
                    "appears_in": row.appears_in_documents or ["not cited"],
                    "in_reference_list": row.in_reference_list,
                    "title_match": row.title_match.value,
                    "author_match": row.author_match.value,
                    "year_match": row.year_match.value,
                    "journal_match": row.journal_match.value,
                    "doi_resolves": row.doi_resolves.value,
                    "outcome": row.outcome.value,
                    "notes": row.notes,
                }
                for index, row in enumerate(self.state.audit_rows, 1)
            ],
            self.job.paths.verification / "Citation_Audit.xlsx",
            self.settings,
        )
        self.publish([evidence_path, audit_path], "verification")
        self.publish([self.job.paths.evidence_file], "logs")

        artefacts = [excel_path, *word_paths, evidence_path, audit_path]
        return StageOutcome(
            StageName.REPORT,
            StageStatus.COMPLETE,
            message=f"{len(artefacts)} report artefacts written.",
            artefacts=artefacts,
            counters={
                "word_documents": len(word_paths),
                "excel_workbooks": 3,
                "evidence_records": len(self.state.ledger.records),
            },
        )

    def _topic_columns(self) -> list[TopicColumn]:
        """Propose topic-specific Master Matrix columns worth adding.

        Only columns that materially help answer the stated research question are
        proposed, and each carries the rationale recorded in the Data Dictionary.
        """
        columns: list[TopicColumn] = []
        analyses = self.state.analyses
        if not analyses:
            return columns

        limit = int(self.settings.reporting.get("max_topic_specific_columns", 10))

        # Detected methods, countries, and software are topic-shaped facets that
        # the standard columns do not expose as filterable lists.
        if any(a.detected_methods for a in analyses.values()):
            columns.append(
                TopicColumn(
                    header="Methods detected in full text",
                    definition=(
                        "Every recognised method named in the paper's own text, as a "
                        "filterable list."
                    ),
                    rationale=(
                        "Lets the reviewer filter the matrix by method, which the single "
                        "'Model / Method' summary column cannot support."
                    ),
                    values={
                        rid: "; ".join(a.detected_methods) or "None detected"
                        for rid, a in analyses.items()
                    },
                )
            )
        if any(a.detected_countries for a in analyses.values()):
            columns.append(
                TopicColumn(
                    header="Study locations detected",
                    definition="Countries and cities named in the paper's text.",
                    rationale=(
                        "Supports the geographic-coverage question the review asks, and "
                        "feeds the landscape statistics."
                    ),
                    values={
                        rid: "; ".join(dict.fromkeys(a.detected_countries)) or "None detected"
                        for rid, a in analyses.items()
                    },
                )
            )
        if any(a.detected_software for a in analyses.values()):
            columns.append(
                TopicColumn(
                    header="Software detected",
                    definition="Software and tools named in the paper's methods.",
                    rationale=(
                        "Identifies reproducible toolchains, which matters for anyone "
                        "planning to replicate this evidence base."
                    ),
                    values={
                        rid: "; ".join(a.detected_software) or "None detected"
                        for rid, a in analyses.items()
                    },
                )
            )
        columns.append(
            TopicColumn(
                header="Evidence stance summary",
                definition=(
                    "How many analysed fields are author-stated versus agent-inferred "
                    "for this paper."
                ),
                rationale=(
                    "Shows at a glance how much of this row rests on the authors' own "
                    "words rather than the agent's reading."
                ),
                width=24,
                values={rid: _stance_summary(a) for rid, a in analyses.items()},
            )
        )
        return columns[:limit]

    # ------------------------------------------------------------------
    # Stage 11: final verification
    # ------------------------------------------------------------------

    def _stage_final_verify(self) -> StageOutcome:
        """Run all three verifiers and write the verification report."""
        if not self.state.records:
            return StageOutcome(
                StageName.FINAL_VERIFY,
                StageStatus.SKIPPED,
                message="No records to verify.",
            )

        included = [r for r in self.state.records if r.selected] or list(self.state.records)
        manager = CitationManager(style=self.job.config.citation_style, records=included)
        document_texts = _read_document_texts(
            list(self.job.paths.reports.glob("*.docx"))
        )

        result = run_verification(
            included,
            self.state.analyses,
            self.state.ledger,
            self.state.audit_rows,
            self.job.config,
            self.settings,
            merge_events=self.state.merge_events,
            document_texts=document_texts,
            discovered_count=self.state.discovered_count,
        )
        write_json(
            self.job.paths.findings_file,
            [f.model_dump(mode="json") for f in result.findings],
        )

        synthesis = build_synthesis(
            included, self.state.analyses, manager, EvidenceLedger(), self.job.config
        )
        context = ReportContext(
            config=self.job.config,
            settings=self.settings,
            records=included,
            analyses=self.state.analyses,
            manager=manager,
            ledger=self.state.ledger,
            synthesis=synthesis,
        )
        report_path = build_verification_report(
            result.summary,
            result.findings,
            context,
            self.job.paths.verification / "Verification_Report.docx",
        )
        issues_path = write_unresolved_issues_csv(
            result, self.job.paths.verification / "unresolved_issues.csv"
        )
        self.publish([report_path, issues_path], "verification")

        self.save_records()
        self.publish([self.job.paths.records_file, self.job.paths.findings_file], "logs")
        self.publish([self.job.paths.job_config_file, self.job.paths.pipeline_log], "logs")

        return StageOutcome(
            StageName.FINAL_VERIFY,
            StageStatus.COMPLETE,
            message=(
                f"{result.summary.passed_checks} checks passed, "
                f"{result.summary.warnings} warning(s), "
                f"{result.summary.unresolved_problems} unresolved. "
                f"Confidence: {result.summary.overall_confidence}."
            ),
            artefacts=[report_path, issues_path],
            counters={
                "passed": result.summary.passed_checks,
                "warnings": result.summary.warnings,
                "unresolved": result.summary.unresolved_problems,
            },
        )

    # ------------------------------------------------------------------
    # Completion summary
    # ------------------------------------------------------------------

    def completion_summary(self) -> dict[str, Any]:
        """Build the final summary shown to the user."""
        records = self.state.records
        selected = [r for r in records if r.selected]
        downloaded = [
            r for r in records
            if r.download_status in (DownloadStatus.DOWNLOADED, DownloadStatus.ALREADY_PRESENT)
        ]
        findings = read_json(self.job.paths.findings_file, []) or []
        unresolved = [f for f in findings if f.get("outcome") == "Fail"]

        drive = self.storage.summary()
        limitations: list[str] = []
        if not self.job.config.ranking_file:
            limitations.append(
                "No licensed journal-ranking file was configured, so quartiles are "
                "reported as Unverified rather than guessed."
            )
        failed = len(selected) - len(downloaded)
        if failed > 0:
            limitations.append(
                f"{failed} selected paper(s) could not be retrieved legally and are "
                "listed in Unable_to_Download.docx."
            )
        if not self.state.analyses:
            limitations.append(
                "No paper was analysed, so the synthesis documents contain no findings."
            )
        if unresolved:
            limitations.append(
                f"{len(unresolved)} verification problem(s) remain unresolved; see "
                "unresolved_issues.csv."
            )
        if not llm_available(self.settings):
            limitations.append(
                f"Claude assistance was not used ({unavailable_reason(self.settings)}). "
                "The deterministic analysis path was used throughout."
            )
        if drive["drive_enabled"] and drive["artefacts_pending_upload"]:
            limitations.append(
                f"{drive['artefacts_pending_upload']} artefact(s) are staged locally but "
                "not yet verified in Google Drive. Run 'drive-sync' to retry."
            )
        elif not drive["drive_enabled"]:
            limitations.append(
                f"Google Drive syncing is not active ({drive['drive_status']}). All "
                "outputs remain in local staging only."
            )

        return {
            "topic": self.job.config.topic,
            "job_folder": str(self.job.paths.logs),
            "job_date": self.job.config.job_date,
            "assumptions": self.job.config.assumptions,
            "counts": {
                "discovered": self.state.discovered_count,
                "unique": len(records),
                "included": len(selected),
                "verified_q1": sum(
                    1 for r in records if r.q1.verification_status == Q1Status.VERIFIED_Q1
                ),
                "unverified_quartile": sum(
                    1 for r in records if r.q1.verification_status == Q1Status.UNVERIFIED
                ),
                "downloaded": len(downloaded),
                "failed_downloads": failed if failed > 0 else 0,
                "analysed": len(self.state.analyses),
                "evidence_records": len(self.state.ledger.records),
                "unresolved_problems": len(unresolved),
            },
            "drive": drive,
            "missing_api_keys": missing_keys(),
            "limitations": limitations,
            "stages": self.job.describe()["stages"],
            "complete": self.job.next_stage() is None,
        }


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _stance_summary(analysis: PaperAnalysis) -> str:
    """Summarise how much of one analysis is author-stated versus inferred."""
    from .schemas import ANALYSIS_FIELDS, EvidenceStance

    author = sum(
        1
        for name in ANALYSIS_FIELDS
        if analysis.field(name).stance == EvidenceStance.AUTHOR_STATED
        and analysis.field(name).is_reported
    )
    inferred = sum(
        1
        for name in ANALYSIS_FIELDS
        if analysis.field(name).stance == EvidenceStance.AGENT_INFERENCE
        and analysis.field(name).is_reported
    )
    unreported = len(ANALYSIS_FIELDS) - author - inferred
    return f"{author} author-stated; {inferred} inferred; {unreported} not reported"


def _read_document_texts(paths: list[Path]) -> dict[str, str]:
    """Read the plain text of generated Word documents for citation scanning."""
    from docx import Document

    texts: dict[str, str] = {}
    for path in paths:
        path = Path(path)
        if not path.exists() or path.suffix.lower() != ".docx":
            continue
        try:
            document = Document(path)
        except Exception as exc:  # noqa: BLE001 - a bad file is skipped, not fatal
            LOG.debug(f"Could not read {path.name} for citation scanning: {exc}")
            continue
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        texts[path.name] = "\n".join(parts)
    return texts


def _read_jsonl(path: Path) -> list[dict[str, Any]]:
    """Read a JSONL log into a list."""
    from .utils import read_jsonl

    return list(read_jsonl(Path(path)))


def create_and_run(
    topic: str,
    *,
    settings: Settings | None = None,
    enable_drive: bool | None = None,
    stages: list[StageName] | None = None,
    **job_overrides: Any,
) -> tuple[Job, Orchestrator, list[StageOutcome]]:
    """Create a job and run the pipeline. The entry point used by the CLI."""
    settings = settings or load_settings()
    job = Job.create(topic, settings=settings, **job_overrides)
    orchestrator = Orchestrator(job, settings=settings, enable_drive=enable_drive)
    outcomes = orchestrator.run(stages=stages)
    return job, orchestrator, outcomes
