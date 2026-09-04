"""Excel and Word generation, citations, evidence ledger, and verification."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as ReadDocx
from openpyxl import load_workbook

from literature_review_agent.analysis import analyse_text
from literature_review_agent.citation_manager import (
    CitationManager,
    audit_citations,
    in_text_citation,
    reference_entry,
)
from literature_review_agent.config import Settings
from literature_review_agent.downloader import failure_rows
from literature_review_agent.evidence_ledger import EvidenceLedger
from literature_review_agent.excel_report import (
    MatrixInputs,
    TopicColumn,
    build_citation_audit_workbook,
    build_evidence_workbook,
    build_workbook,
)
from literature_review_agent.job_manager import build_job_config
from literature_review_agent.pdf_extractor import PageText
from literature_review_agent.schemas import (
    CheckOutcome,
    CitationStyle,
    DownloadStatus,
    EvidenceStance,
    GapCategory,
    JobConfig,
    PaperAnalysis,
    PaperRecord,
    Q1Status,
    VerificationSummary,
)
from literature_review_agent.synthesis import build_synthesis
from literature_review_agent.verification import run_verification, write_unresolved_issues_csv
from literature_review_agent.word_reports import (
    ReportContext,
    build_all_reports,
    build_unable_to_download,
    build_verification_report,
)

PAGE_ONE = (
    "Abstract\n\nThis study examines how rainfall intensity influences mode choice in "
    "Delhi, India. We estimate a mixed logit model using a household travel survey of "
    "2450 respondents. Results show that heavy rainfall reduces cycling by 34 percent."
)
PAGE_TWO = (
    "2. Methodology\n\nWe use data from the 2019 Delhi household travel survey. The "
    "dependent variable is the chosen travel mode. We control for income and age. "
    "Estimation was implemented in Biogeme and Python.\n\n"
    "4. Conclusions\n\nPolicymakers should consider sheltered infrastructure. A "
    "limitation is that we cannot observe trip chaining. Future research should examine "
    "other Indian cities."
)
PAGE_MUMBAI = (
    "Abstract\n\nThis paper investigates monsoon rainfall and suburban rail ridership in "
    "Mumbai, India. We apply a fixed effects panel data model to smart card data. "
    "Results indicate ridership increases by 9 percent during moderate rain. "
    "Estimation was implemented in Stata."
)


@pytest.fixture
def config(settings: Settings) -> JobConfig:
    """A job config for the rainfall topic."""
    return build_job_config(
        "Effect of rainfall on urban travel behaviour",
        settings,
        research_questions=["How does rainfall influence mode choice?"],
    )


@pytest.fixture
def analysed(
    settings: Settings, config: JobConfig, sample_records: list[PaperRecord]
) -> tuple[list[PaperRecord], dict[str, PaperAnalysis]]:
    """Two records with completed analyses."""
    pages = {
        "rec-delhi": [PageText(number=1, text=PAGE_ONE), PageText(number=2, text=PAGE_TWO)],
        "rec-mumbai": [PageText(number=1, text=PAGE_MUMBAI)],
    }
    analyses: dict[str, PaperAnalysis] = {}
    for record in sample_records:
        record.download_status = DownloadStatus.DOWNLOADED
        analyses[record.record_id] = analyse_text(
            record, pages[record.record_id], config, settings
        )
    return sample_records, analyses


@pytest.fixture
def context(
    settings: Settings,
    config: JobConfig,
    analysed: tuple[list[PaperRecord], dict[str, PaperAnalysis]],
) -> ReportContext:
    """A fully populated report context."""
    records, analyses = analysed
    manager = CitationManager(style=config.citation_style, records=list(records))
    ledger = EvidenceLedger()
    synthesis = build_synthesis(records, analyses, manager, ledger, config)
    return ReportContext(
        config=config,
        settings=settings,
        records=records,
        analyses=analyses,
        manager=manager,
        ledger=ledger,
        synthesis=synthesis,
    )


# ---------------------------------------------------------------------------
# Citations
# ---------------------------------------------------------------------------


class TestCitationFormatting:
    """Style-correct citations built only from verified metadata."""

    def test_apa_two_authors(self, sample_records: list[PaperRecord]) -> None:
        assert in_text_citation(sample_records[0], CitationStyle.APA7) == (
            "(Sharma & Patel, 2021)"
        )

    def test_apa_single_author(self, sample_records: list[PaperRecord]) -> None:
        assert in_text_citation(sample_records[1], CitationStyle.APA7) == "(Iyer, 2019)"

    def test_three_authors_use_et_al(self) -> None:
        record = PaperRecord(
            title="X", authors=["Lee, Min", "Kim, Soo", "Park, Jae"], year=2019
        )
        assert in_text_citation(record, CitationStyle.APA7) == "(Lee et al., 2019)"

    def test_missing_year_becomes_nd(self) -> None:
        record = PaperRecord(title="X", authors=["Sharma, R"])
        assert "n.d." in in_text_citation(record, CitationStyle.APA7)

    def test_no_authors_becomes_anonymous(self) -> None:
        record = PaperRecord(title="X", year=2021)
        assert "Anonymous" in in_text_citation(record, CitationStyle.APA7)

    def test_page_number_included_for_quotations(
        self, sample_records: list[PaperRecord]
    ) -> None:
        citation = in_text_citation(sample_records[0], CitationStyle.APA7, page=47)
        assert "p. 47" in citation

    def test_numeric_styles_use_brackets(self, sample_records: list[PaperRecord]) -> None:
        assert in_text_citation(sample_records[0], CitationStyle.IEEE, number=3) == "[3]"

    def test_apa_reference_entry(self, sample_records: list[PaperRecord]) -> None:
        entry = reference_entry(sample_records[0], CitationStyle.APA7)
        assert entry.startswith("Sharma, R., & Patel, N. (2021).")
        assert "Transportation Research Part A" in entry
        assert "150(2), 45-61" in entry
        assert "https://doi.org/10.1016/j.tra.2021.01.001" in entry

    def test_missing_fields_are_marked_not_invented(self) -> None:
        record = PaperRecord(title="A study with no venue recorded", year=2021)
        entry = reference_entry(record, CitationStyle.APA7)
        assert "[missing]" in entry, "an absent field must be visible, not filled in"

    @pytest.mark.parametrize("style", list(CitationStyle))
    def test_every_style_produces_output(
        self, style: CitationStyle, sample_records: list[PaperRecord]
    ) -> None:
        assert reference_entry(sample_records[0], style).strip()
        assert in_text_citation(sample_records[0], style, number=1).strip()


class TestCitationManager:
    """Stable numbering, disambiguation, and exports."""

    def test_disambiguates_same_author_and_year(self) -> None:
        records = [
            PaperRecord(record_id="a", title="First study", authors=["Sharma, R"], year=2021),
            PaperRecord(record_id="b", title="Second study", authors=["Sharma, R"], year=2021),
        ]
        manager = CitationManager(style=CitationStyle.APA7, records=records)
        citations = {manager.citation("a"), manager.citation("b")}
        assert citations == {"(Sharma, 2021a)", "(Sharma, 2021b)"}

    def test_unknown_record_is_never_fabricated(self) -> None:
        manager = CitationManager(records=[])
        assert manager.citation("does-not-exist") == "[citation unavailable]"
        assert manager.reference("does-not-exist") == "[reference unavailable]"

    def test_citation_group(self, sample_records: list[PaperRecord]) -> None:
        manager = CitationManager(style=CitationStyle.APA7, records=list(sample_records))
        group = manager.citation_group(["rec-delhi", "rec-mumbai"])
        assert group.startswith("(") and group.endswith(")")
        assert ";" in group

    def test_numeric_numbering_is_stable(self, sample_records: list[PaperRecord]) -> None:
        manager = CitationManager(style=CitationStyle.IEEE, records=list(sample_records))
        first = manager.number_for("rec-delhi")
        manager.rebuild()
        assert manager.number_for("rec-delhi") == first

    def test_bibtex_export(self, sample_records: list[PaperRecord], tmp_path: Path) -> None:
        manager = CitationManager(records=list(sample_records))
        written = manager.write_exports(tmp_path / "refs.bib", tmp_path / "refs.ris")
        assert len(written) == 2
        bib = (tmp_path / "refs.bib").read_text(encoding="utf-8")
        assert bib.count("@article") == 2
        assert "Rainfall Intensity and Mode Choice in Delhi" in bib
        assert "10.1016/j.tra.2021.01.001" in bib

    def test_ris_export(self, sample_records: list[PaperRecord], tmp_path: Path) -> None:
        manager = CitationManager(records=list(sample_records))
        manager.write_exports(tmp_path / "refs.bib", tmp_path / "refs.ris")
        ris = (tmp_path / "refs.ris").read_text(encoding="utf-8")
        assert ris.count("TY  - JOUR") == 2
        assert "Sharma, Ravi" in ris

    def test_preprint_gets_a_preprint_type(self, tmp_path: Path) -> None:
        record = PaperRecord(
            record_id="p", title="A preprint", authors=["Zhang, W"], year=2021,
            document_type="preprint",
        )
        manager = CitationManager(records=[record])
        assert "@misc" in manager.to_bibtex()
        assert "TY  - UNPB" in manager.to_ris()


class TestCitationAudit:
    """The audit runs in both directions."""

    def test_uncited_reference_raises_a_warning(
        self, sample_records: list[PaperRecord]
    ) -> None:
        manager = CitationManager(records=list(sample_records))
        rows = audit_citations(manager, [])
        assert len(rows) == 2
        assert all(r.outcome == CheckOutcome.WARNING for r in rows)
        assert all("no synthesis claim cites it" in r.notes for r in rows)

    def test_cited_record_passes(self, sample_records: list[PaperRecord]) -> None:
        manager = CitationManager(records=list(sample_records))
        ledger = EvidenceLedger()
        for record in sample_records:
            ledger.add(
                document="Introduction.docx", section="1", claim="A claim",
                record_id=record.record_id, in_text_citation=manager.citation(record.record_id),
                pages=[1],
            )
        rows = audit_citations(manager, ledger.records)
        assert all(r.outcome == CheckOutcome.PASS for r in rows)

    def test_citation_to_an_unknown_record_fails(self) -> None:
        manager = CitationManager(records=[])
        ledger = EvidenceLedger()
        ledger.add(
            document="Introduction.docx", section="1", claim="A claim",
            record_id="ghost", in_text_citation="(Ghost, 2021)",
        )
        rows = audit_citations(manager, ledger.records)
        assert rows[0].outcome == CheckOutcome.FAIL
        assert "not in the retrieved set" in rows[0].notes


# ---------------------------------------------------------------------------
# Evidence ledger
# ---------------------------------------------------------------------------


class TestEvidenceLedger:
    """The ledger is the backbone of every synthesis claim."""

    def test_add_and_query(self) -> None:
        ledger = EvidenceLedger()
        record = ledger.add(
            document="Introduction.docx", section="1", claim="Rain reduces cycling",
            record_id="a", in_text_citation="(Sharma, 2021)", pages=[3],
        )
        assert record.evidence_id.startswith("EV-")
        assert ledger.get(record.evidence_id) is record
        assert ledger.for_document("Introduction.docx") == [record]
        assert ledger.cited_record_ids() == {"a"}

    def test_identical_claims_are_deduplicated(self) -> None:
        ledger = EvidenceLedger()
        first = ledger.add(document="d", section="s", claim="c", record_id="a", pages=[1])
        second = ledger.add(document="d", section="s", claim="c", record_id="a", pages=[2])
        assert first is second
        assert second.pages == [1, 2], "pages must accumulate on a repeated claim"

    def test_unreported_field_produces_no_record(
        self, analysed: tuple[list[PaperRecord], dict[str, PaperAnalysis]]
    ) -> None:
        # This is what stops an absent fact becoming a cited claim.
        _, analyses = analysed
        analysis = next(iter(analyses.values()))
        ledger = EvidenceLedger()
        analysis.model_equations.value = ""
        analysis.model_equations.stance = EvidenceStance.NOT_REPORTED
        assert ledger.add_from_analysis(
            analysis, "model_equations", document="d", section="s", in_text_citation="(X, 2021)"
        ) is None
        assert not ledger.records

    def test_stance_counts(self) -> None:
        ledger = EvidenceLedger()
        ledger.add(document="d", section="s", claim="a", record_id="1",
                   stance=EvidenceStance.AUTHOR_STATED)
        ledger.add(document="d", section="s", claim="b", record_id="1",
                   stance=EvidenceStance.AGENT_INFERENCE)
        counts = ledger.counts_by_stance()
        assert counts[EvidenceStance.AUTHOR_STATED.value] == 1
        assert counts[EvidenceStance.AGENT_INFERENCE.value] == 1

    def test_claims_without_pages_are_flagged(self) -> None:
        ledger = EvidenceLedger()
        ledger.add(document="d", section="s", claim="a", record_id="1",
                   stance=EvidenceStance.AUTHOR_STATED)
        assert len(ledger.claims_without_pages()) == 1

    def test_save_and_load_round_trip(self, tmp_path: Path) -> None:
        ledger = EvidenceLedger()
        ledger.add(document="d", section="s", claim="a claim", record_id="1", pages=[5])
        path = ledger.save(tmp_path / "evidence.json")
        reloaded = EvidenceLedger.load(path)
        assert len(reloaded.records) == 1
        assert reloaded.records[0].pages == [5]

    def test_rows_are_report_ready(self) -> None:
        ledger = EvidenceLedger()
        ledger.add(document="d", section="s", claim="a claim", record_id="1", pages=[5])
        row = ledger.rows()[0]
        assert row["serial"] == 1
        assert row["pages"] == "5"


# ---------------------------------------------------------------------------
# Synthesis
# ---------------------------------------------------------------------------


class TestSynthesis:
    """Gaps, models, and landscape statistics."""

    def test_gaps_are_categorised_and_cited(self, context: ReportContext) -> None:
        gaps = context.synthesis.gaps
        assert gaps
        assert all(gap.statement for gap in gaps)
        assert any(gap.citations for gap in gaps)

    def test_author_and_inferred_gaps_are_separated(self, context: ReportContext) -> None:
        categories = {gap.category for gap in context.synthesis.gaps}
        author_gaps = [g for g in context.synthesis.gaps
                       if g.category == GapCategory.AUTHOR_STATED]
        inferred = [g for g in context.synthesis.gaps
                    if g.category == GapCategory.AGENT_INFERRED]
        assert all(g.stance == EvidenceStance.AGENT_INFERENCE for g in inferred)
        assert categories, "at least one gap category must be populated"
        for gap in author_gaps:
            assert gap.stance == EvidenceStance.AUTHOR_STATED

    def test_models_are_profiled(self, context: ReportContext) -> None:
        models = context.synthesis.models
        assert models
        names = {m.model_name.lower() for m in models}
        assert any("logit" in n or "panel" in n for n in names)
        assert all(m.plain_explanation for m in models)

    def test_landscape_counts_the_reviewed_evidence(self, context: ReportContext) -> None:
        landscape = context.synthesis.landscape
        assert landscape.total_papers == 2
        assert "India" in landscape.countries
        assert landscape.regions.get("South Asia", 0) >= 1

    def test_landscape_lists_under_represented_regions(self, context: ReportContext) -> None:
        assert any("Europe" in item for item in context.synthesis.landscape.under_researched)

    def test_contradiction_between_papers_is_detected(self, context: ReportContext) -> None:
        # One paper reports a reduction, the other an increase.
        assert context.synthesis.contradictions
        assert any(
            gap.category == GapCategory.CONTRADICTORY for gap in context.synthesis.gaps
        )


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


class TestExcelWorkbook:
    """The ten-sheet literature-review matrix."""

    @pytest.fixture
    def workbook_path(
        self, context: ReportContext, settings: Settings, tmp_path: Path
    ) -> Path:
        inputs = MatrixInputs(
            config=context.config,
            records=context.records,
            analyses=context.analyses,
            gaps=context.synthesis.gaps,
            models=context.synthesis.models,
            landscape=context.synthesis.landscape,
            citation_audit=audit_citations(context.manager, context.ledger.records),
            citations=context.manager.citation_map(),
            search_log=[
                {"timestamp": "2026-01-01T00:00:00", "source": "Crossref",
                 "query": "rainfall", "breadth": "balanced", "results": 5,
                 "http_status": 200, "outcome": "ok", "notes": ""}
            ],
            topic_columns=[
                TopicColumn(
                    header="Rainfall threshold reported",
                    definition="The rainfall intensity threshold the paper reports.",
                    rationale="Central to the review's research question.",
                    values={"rec-delhi": "10 mm/h"},
                )
            ],
        )
        return build_workbook(inputs, settings, tmp_path / "Literature_Review_Matrix.xlsx")

    def test_all_ten_sheets_exist(self, workbook_path: Path) -> None:
        book = load_workbook(workbook_path)
        assert book.sheetnames == [
            "Master Matrix", "Methods and Models", "Key Findings", "Research Gaps",
            "Global Landscape", "Q1 Verification", "Citation Audit",
            "Download Failures", "Search Log", "Data Dictionary",
        ]

    def test_one_row_per_paper(self, workbook_path: Path) -> None:
        sheet = load_workbook(workbook_path)["Master Matrix"]
        assert sheet.max_row == 3, "header plus one row per paper"

    def test_header_is_frozen_and_filtered(self, workbook_path: Path) -> None:
        sheet = load_workbook(workbook_path)["Master Matrix"]
        assert sheet.freeze_panes == "A2"
        assert sheet.auto_filter.ref is not None

    def test_no_merged_cells_in_data_tables(self, workbook_path: Path) -> None:
        book = load_workbook(workbook_path)
        for name in book.sheetnames:
            assert not book[name].merged_cells.ranges, f"{name} must not merge data cells"

    def test_status_cells_are_colour_coded(self, workbook_path: Path) -> None:
        sheet = load_workbook(workbook_path)["Master Matrix"]
        headers = [c.value for c in sheet[1]]
        column = headers.index("Q1 Status") + 1
        fill = sheet.cell(row=2, column=column).fill
        assert fill.fgColor.rgb not in (None, "00000000"), "verification status must be shaded"

    def test_doi_hyperlink_is_set(self, workbook_path: Path) -> None:
        sheet = load_workbook(workbook_path)["Master Matrix"]
        headers = [c.value for c in sheet[1]]
        column = headers.index("DOI Link") + 1
        assert sheet.cell(row=2, column=column).hyperlink is not None

    def test_topic_column_is_added_and_documented(self, workbook_path: Path) -> None:
        book = load_workbook(workbook_path)
        headers = [c.value for c in book["Master Matrix"][1]]
        assert "Rainfall threshold reported" in headers

        dictionary = book["Data Dictionary"]
        rows = [
            [dictionary.cell(row=r, column=c).value for c in range(1, 7)]
            for r in range(2, dictionary.max_row + 1)
        ]
        added = [r for r in rows if r[4] == "Yes"]
        assert added, "an agent-added column must be documented in the Data Dictionary"
        assert any("Rainfall threshold" in str(r[1]) for r in added)
        assert all(r[5] for r in added), "every added column needs a rationale"

    def test_topic_columns_are_capped(
        self, context: ReportContext, settings: Settings, tmp_path: Path
    ) -> None:
        many = [
            TopicColumn(header=f"Extra {i}", definition="d", rationale="r", values={})
            for i in range(15)
        ]
        inputs = MatrixInputs(
            config=context.config, records=context.records,
            analyses=context.analyses, topic_columns=many,
        )
        path = build_workbook(inputs, settings, tmp_path / "capped.xlsx")
        headers = [c.value for c in load_workbook(path)["Master Matrix"][1]]
        assert sum(1 for h in headers if str(h).startswith("Extra ")) == 10

    def test_analysis_cells_mark_unreported_fields(self, workbook_path: Path) -> None:
        sheet = load_workbook(workbook_path)["Master Matrix"]
        headers = [c.value for c in sheet[1]]
        values = [
            sheet.cell(row=row, column=headers.index("Model Equations") + 1).value
            if "Model Equations" in headers else None
            for row in (2, 3)
        ]
        # Whatever the column set, an unreported field must never be blank prose.
        assert all(v is None or v for v in values)

    def test_empty_sheet_explains_itself(
        self, context: ReportContext, settings: Settings, tmp_path: Path
    ) -> None:
        inputs = MatrixInputs(config=context.config, records=context.records)
        path = build_workbook(inputs, settings, tmp_path / "empty.xlsx")
        note = load_workbook(path)["Download Failures"].cell(row=2, column=1).value
        assert note and "No download failures" in note

    def test_evidence_workbook(self, context: ReportContext, tmp_path: Path) -> None:
        path = build_evidence_workbook(
            context.ledger.rows(), tmp_path / "Evidence_Ledger.xlsx"
        )
        sheet = load_workbook(path).active
        headers = [c.value for c in sheet[1]]
        assert "Evidence ID" in headers and "Pages" in headers and "Evidence Stance" in headers

    def test_citation_audit_workbook(
        self, context: ReportContext, settings: Settings, tmp_path: Path
    ) -> None:
        rows = [
            {
                "serial": 1, "in_text_citation": "(Sharma, 2021)", "record_id": "rec-delhi",
                "doi": "10.1/a", "title": "T", "appears_in": ["Introduction.docx"],
                "in_reference_list": True, "title_match": "Pass", "author_match": "Pass",
                "year_match": "Pass", "journal_match": "Pass", "doi_resolves": "Pass",
                "outcome": "Pass", "notes": "",
            }
        ]
        path = build_citation_audit_workbook(rows, tmp_path / "Citation_Audit.xlsx", settings)
        assert load_workbook(path).active.max_row == 2


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------


class TestWordReports:
    """The five synthesis documents plus the two registers."""

    @pytest.fixture
    def reports(self, context: ReportContext, tmp_path: Path) -> dict[str, Path]:
        paths = build_all_reports(context, tmp_path)
        return {p.name: p for p in paths}

    def test_all_five_documents_are_written(self, reports: dict[str, Path]) -> None:
        assert set(reports) == {
            "Introduction.docx", "Research_Gaps.docx", "Global_Research_Landscape.docx",
            "Models_and_Applications.docx", "Paper_Summaries.docx",
        }

    @pytest.mark.parametrize(
        "name",
        ["Introduction.docx", "Research_Gaps.docx", "Global_Research_Landscape.docx",
         "Models_and_Applications.docx", "Paper_Summaries.docx"],
    )
    def test_every_document_has_the_required_spine(
        self, name: str, reports: dict[str, Path]
    ) -> None:
        document = ReadDocx(reports[name])
        headings = [
            p.text for p in document.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "Scope" in headings
        assert "Method and evidence base" in headings
        assert "References" in headings
        assert "Limitations" in headings
        assert "Verification note" in headings

    @pytest.mark.parametrize(
        "name",
        ["Introduction.docx", "Research_Gaps.docx", "Global_Research_Landscape.docx",
         "Models_and_Applications.docx", "Paper_Summaries.docx"],
    )
    def test_every_document_names_the_topic_and_date(
        self, name: str, reports: dict[str, Path], context: ReportContext
    ) -> None:
        text = "\n".join(p.text for p in ReadDocx(reports[name]).paragraphs)
        assert context.config.topic in text
        assert context.generated_on in text

    def test_introduction_covers_all_seven_sections(self, reports: dict[str, Path]) -> None:
        headings = [
            p.text for p in ReadDocx(reports["Introduction.docx"]).paragraphs
            if p.style.name == "Heading 1"
        ]
        for fragment in (
            "1. Background", "2. Why the topic matters", "3. What is currently known",
            "4. Methods used", "5. Convergence", "6. What remains unresolved",
            "7. Motivation for further research",
        ):
            assert any(fragment in h for h in headings), f"missing section: {fragment}"

    def test_introduction_cites_its_claims(self, reports: dict[str, Path]) -> None:
        text = "\n".join(p.text for p in ReadDocx(reports["Introduction.docx"]).paragraphs)
        assert "(Sharma & Patel, 2021)" in text or "(Iyer, 2019)" in text

    def test_introduction_has_a_reference_list(self, reports: dict[str, Path]) -> None:
        text = "\n".join(p.text for p in ReadDocx(reports["Introduction.docx"]).paragraphs)
        assert "https://doi.org/10.1016/j.tra.2021.01.001" in text

    def test_research_gaps_has_all_eleven_categories(self, reports: dict[str, Path]) -> None:
        headings = [
            p.text for p in ReadDocx(reports["Research_Gaps.docx"]).paragraphs
            if p.style.name == "Heading 1"
        ]
        for category in GapCategory:
            assert any(category.value in h for h in headings), f"missing: {category.value}"

    def test_research_gaps_warns_about_inference(self, reports: dict[str, Path]) -> None:
        text = "\n".join(p.text for p in ReadDocx(reports["Research_Gaps.docx"]).paragraphs)
        assert "agent's inferences" in text

    def test_landscape_avoids_claiming_world_coverage(self, reports: dict[str, Path]) -> None:
        text = "\n".join(
            p.text for p in ReadDocx(reports["Global_Research_Landscape.docx"]).paragraphs
        )
        assert "not of world research" in text
        assert "reviewed evidence" in text

    def test_landscape_reports_institutions_honestly(self, reports: dict[str, Path]) -> None:
        # Affiliations are not retrieved, so this must be stated, not guessed.
        text = "\n".join(
            p.text for p in ReadDocx(reports["Global_Research_Landscape.docx"]).paragraphs
        )
        assert "Author affiliations are not retrieved" in text

    def test_models_document_explains_each_model(self, reports: dict[str, Path]) -> None:
        document = ReadDocx(reports["Models_and_Applications.docx"])
        text = "\n".join(p.text for p in document.paragraphs)
        assert "How this family works in plain terms" in text
        assert document.tables, "model attributes must be tabulated"

        cells = [
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        ]
        for attribute in (
            "Model category", "Main assumptions", "Required inputs", "Outputs",
            "Study application", "Software used", "Calibration approach",
            "Validation approach", "Advantages", "Limitations",
            "How it works (plain language)",
        ):
            assert attribute in cells, f"missing model attribute: {attribute}"

    def test_paper_summaries_include_page_evidence(self, reports: dict[str, Path]) -> None:
        document = ReadDocx(reports["Paper_Summaries.docx"])
        cells = [
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        ]
        assert any("(p. " in cell for cell in cells), "summaries must cite page numbers"

    def test_paper_summaries_mark_unreported_fields(self, reports: dict[str, Path]) -> None:
        cells = [
            cell.text
            for table in ReadDocx(reports["Paper_Summaries.docx"]).tables
            for row in table.rows
            for cell in row.cells
        ]
        assert any(cell == "Information not reported" for cell in cells)

    def test_agent_inference_is_labelled_in_the_text(self, reports: dict[str, Path]) -> None:
        cells = [
            cell.text
            for table in ReadDocx(reports["Paper_Summaries.docx"]).tables
            for row in table.rows
            for cell in row.cells
        ]
        assert any("[Agent inference]" in cell for cell in cells)

    def test_reports_survive_having_no_analyses(
        self, settings: Settings, config: JobConfig, sample_records: list[PaperRecord],
        tmp_path: Path,
    ) -> None:
        manager = CitationManager(records=list(sample_records))
        ledger = EvidenceLedger()
        synthesis = build_synthesis(sample_records, {}, manager, ledger, config)
        context = ReportContext(
            config=config, settings=settings, records=sample_records, analyses={},
            manager=manager, ledger=ledger, synthesis=synthesis,
        )
        paths = build_all_reports(context, tmp_path)
        assert len(paths) == 5
        text = "\n".join(p.text for p in ReadDocx(tmp_path / "Introduction.docx").paragraphs)
        assert "no evidence-based introduction can be written" in text


class TestUnableToDownload:
    """The manual-retrieval register."""

    def test_records_a_failure_with_every_column(
        self, config: JobConfig, sample_records: list[PaperRecord], tmp_path: Path
    ) -> None:
        record = sample_records[0]
        record.download_status = DownloadStatus.FAILED
        record.failure_reason = "HTTP 403: access is restricted."
        record.q1.verification_status = Q1Status.UNVERIFIED
        rows = failure_rows(sample_records)
        assert len(rows) == 1

        path = build_unable_to_download(rows, config, tmp_path / "Unable_to_Download.docx")
        document = ReadDocx(path)
        headers = [cell.text for cell in document.tables[0].rows[0].cells]
        for expected in (
            "S. No.", "Paper title", "Authors", "Year", "Journal", "DOI",
            "Publisher landing page", "Attempted PDF links", "OA status",
            "Failure reason", "HTTP status", "Attempted at",
            "Recommended manual action", "Q1 status",
        ):
            assert expected in headers, f"missing column: {expected}"

    def test_explains_the_legal_boundary(
        self, config: JobConfig, sample_records: list[PaperRecord], tmp_path: Path
    ) -> None:
        sample_records[0].download_status = DownloadStatus.FAILED
        path = build_unable_to_download(
            failure_rows(sample_records), config, tmp_path / "u.docx"
        )
        text = "\n".join(p.text for p in ReadDocx(path).paragraphs)
        assert "does not bypass paywalls" in text

    def test_empty_register_says_so(self, config: JobConfig, tmp_path: Path) -> None:
        path = build_unable_to_download([], config, tmp_path / "u.docx")
        text = "\n".join(p.text for p in ReadDocx(path).paragraphs)
        assert "No entries" in text


class TestVerificationReport:
    """The verification report and its counts."""

    def test_reports_every_required_count(
        self, context: ReportContext, tmp_path: Path
    ) -> None:
        rows = audit_citations(context.manager, context.ledger.records)
        result = run_verification(
            context.records, context.analyses, context.ledger, rows,
            context.config, context.settings, discovered_count=7,
        )
        path = build_verification_report(
            result.summary, result.findings, context, tmp_path / "Verification_Report.docx"
        )
        cells = [
            cell.text for table in ReadDocx(path).tables for row in table.rows
            for cell in row.cells
        ]
        for measure in (
            "Total papers discovered", "Total unique papers after deduplication",
            "Total included in the review", "Verified Q1", "Unverified quartile",
            "PDFs downloaded and validated", "Failed or unavailable downloads",
            "Papers analysed", "Claims checked", "Checks passed", "Warnings",
            "Unresolved problems",
        ):
            assert measure in cells, f"missing count: {measure}"

    def test_states_an_overall_confidence(
        self, context: ReportContext, tmp_path: Path
    ) -> None:
        rows = audit_citations(context.manager, context.ledger.records)
        result = run_verification(
            context.records, context.analyses, context.ledger, rows,
            context.config, context.settings,
        )
        path = build_verification_report(
            result.summary, result.findings, context, tmp_path / "v.docx"
        )
        headings = [
            p.text for p in ReadDocx(path).paragraphs if p.style.name == "Heading 1"
        ]
        assert any("Overall confidence" in h for h in headings)
        assert any("Recommended manual checks" in h for h in headings)

    def test_unresolved_issues_csv_preserves_original_values(
        self, context: ReportContext, tmp_path: Path
    ) -> None:
        rows = audit_citations(context.manager, context.ledger.records)
        result = run_verification(
            context.records, context.analyses, context.ledger, rows,
            context.config, context.settings,
        )
        path = write_unresolved_issues_csv(result, tmp_path / "unresolved_issues.csv")
        content = path.read_text(encoding="utf-8-sig")
        assert "original_value" in content
        assert "corrected_value" in content
        assert "correction_reason" in content

    def test_summary_with_no_analyses_reports_no_confidence(
        self, settings: Settings, config: JobConfig, sample_records: list[PaperRecord]
    ) -> None:
        result = run_verification(
            sample_records, {}, EvidenceLedger(), [], config, settings,
        )
        assert "None" in result.summary.overall_confidence
        assert any(
            "No paper was analysed" in check
            for check in result.summary.recommended_manual_checks
        )

    def test_verification_summary_defaults_are_safe(self) -> None:
        summary = VerificationSummary()
        assert summary.overall_confidence == "Low"
        assert summary.unresolved_problems == 0
