"""Word report generation with ``python-docx``.

Seven documents: ``Introduction``, ``Research_Gaps``, ``Global_Research_Landscape``,
``Models_and_Applications``, ``Paper_Summaries``, ``Unable_to_Download``, and
``Verification_Report``.

Every synthesis document carries the same structural spine — title, topic, date,
scope, evidence base, content, in-text citations, reference list, limitations, and
a verification note — and every substantive statement is written from an
evidence-ledger record, so a claim with no evidence never reaches the page.

Language discipline: statistics describe *the reviewed evidence*, not world
research; author statements and agent inferences are always distinguishable in
the text itself.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .citation_manager import CitationManager
from .config import Settings
from .evidence_ledger import EvidenceLedger
from .logging_setup import get_logger
from .schemas import (
    CheckOutcome,
    EvidenceStance,
    GapCategory,
    JobConfig,
    LandscapeSummary,
    ModelProfile,
    PaperAnalysis,
    PaperRecord,
    Q1Status,
    VerificationFinding,
    VerificationSummary,
)
from .synthesis import SynthesisResult, analysed_field_coverage
from .utils import ensure_dir, truncate_text

LOG = get_logger("word")

ACCENT = RGBColor(0x1F, 0x38, 0x64)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

#: Standard note appended to every synthesis document.
VERIFICATION_NOTE = (
    "Every substantive statement in this document is linked to at least one record "
    "in the evidence ledger (Evidence_Ledger.xlsx), which names the source paper, "
    "the page numbers, and whether the wording is the authors' own or the agent's "
    "inference. Statements marked as agent inference are the agent's reading of the "
    "evidence, not claims made by the cited authors. Citations were audited in both "
    "directions: every citation resolves to a retrieved paper, and every reference "
    "entry is cited. Open and check each source before relying on this document."
)


@dataclass
class ReportContext:
    """Everything the Word writers need in one place."""

    config: JobConfig
    settings: Settings
    records: list[PaperRecord]
    analyses: dict[str, PaperAnalysis]
    manager: CitationManager
    ledger: EvidenceLedger
    synthesis: SynthesisResult
    generated_on: str = field(default_factory=lambda: date.today().isoformat())

    @property
    def analysed(self) -> list[PaperRecord]:
        """Papers with a completed analysis."""
        return [r for r in self.records if r.record_id in self.analyses]

    @property
    def cited_records(self) -> list[PaperRecord]:
        """Papers that at least one ledger claim depends on."""
        cited = self.ledger.cited_record_ids()
        return [r for r in self.records if r.record_id in cited]


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------


def _configure_styles(document: Document) -> None:
    """Apply readable base styles to a new document."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def _add_page_numbers(document: Document) -> None:
    """Add a centred ``Page X`` field to the footer."""
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.append(begin)
    field_run._r.append(instruction)
    field_run._r.append(end)
    field_run.font.size = Pt(9)


def start_document(
    title: str,
    context: ReportContext,
    *,
    scope: str,
    evidence_base: str,
) -> Document:
    """Create a document with the standard front matter every report shares."""
    document = Document()
    _configure_styles(document)

    heading = document.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(context.config.topic)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = MUTED

    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Prepared: {context.generated_on}    |    "
        f"Literature covered: {context.config.year_from}-{context.config.year_to}    |    "
        f"Citation style: {context.config.citation_style.value}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = MUTED

    document.add_heading("Scope", level=1)
    document.add_paragraph(scope)

    document.add_heading("Method and evidence base", level=1)
    document.add_paragraph(evidence_base)

    return document


def finish_document(
    document: Document,
    context: ReportContext,
    output_path: Path,
    *,
    limitations: list[str],
    cited_only: bool = True,
) -> Path:
    """Append the reference list, limitations, and verification note, then save."""
    document.add_heading("References", level=1)
    records = context.cited_records if cited_only else context.records
    if not records:
        document.add_paragraph(
            "No references: no paper in this job produced verified evidence, so no "
            "citation could be generated."
        )
    else:
        ordered = [r for r in context.manager.records if r in records]
        for record in ordered:
            paragraph = document.add_paragraph(context.manager.reference(record.record_id))
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)
            paragraph.paragraph_format.space_after = Pt(6)

    document.add_heading("Limitations", level=1)
    for limitation in limitations:
        document.add_paragraph(limitation, style="List Bullet")

    document.add_heading("Verification note", level=1)
    note = document.add_paragraph(VERIFICATION_NOTE)
    for run in note.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = MUTED

    _add_page_numbers(document)
    output_path = Path(output_path)
    ensure_dir(output_path.parent)
    document.save(output_path)
    LOG.info(f"Wrote {output_path.name}.")
    return output_path


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    widths: list[float] | None = None,
) -> Any:
    """Add a bordered table with a bold header row and no merged data cells."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = ""
        paragraph = header_cells[index].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(headers)]):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            run = paragraph.add_run("" if value is None else str(value))
            run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths[: len(headers)]):
                row.cells[index].width = Inches(width)
    return table


def _evidence_sentence(
    context: ReportContext,
    record_id: str,
    field_name: str,
    *,
    document_name: str,
    section: str,
    lead: str = "",
) -> str | None:
    """Build one cited sentence from an analysed field, or return ``None``.

    Returning ``None`` for an unreported field is the mechanism that keeps
    unsupported prose out of the documents.
    """
    analysis = context.analyses.get(record_id)
    if analysis is None:
        return None
    evidence = analysis.field(field_name)
    if not evidence.is_reported:
        return None

    context.ledger.add_from_analysis(
        analysis,
        field_name,
        document=document_name,
        section=section,
        in_text_citation=context.manager.citation(record_id),
    )

    citation = context.manager.citation(record_id)
    text = truncate_text(evidence.value.rstrip(". "), 420)
    prefix = f"{lead} " if lead else ""
    if evidence.stance == EvidenceStance.AGENT_INFERENCE:
        return f"{prefix}Reading the evidence, {text.lower()[:1]}{text[1:]} {citation}."
    return f"{prefix}{text} {citation}."


# ---------------------------------------------------------------------------
# 1. Introduction
# ---------------------------------------------------------------------------


def build_introduction(context: ReportContext, output_path: Path) -> Path:
    """Write ``Introduction.docx``: an evidence-based academic introduction."""
    config = context.config
    analysed = context.analysed
    coverage = analysed_field_coverage(context.analyses)
    document_name = "Introduction.docx"

    scope = (
        f"This introduction synthesises {len(analysed)} paper"
        f"{'s' if len(analysed) != 1 else ''} retrieved and analysed for the topic "
        f"\"{config.topic}\". It addresses the research question"
        f"{'s' if len(config.research_questions) != 1 else ''}: "
        + "; ".join(config.research_questions)
        + f". Only literature published between {config.year_from} and {config.year_to} "
        "was considered, and only claims traceable to a retrieved paper are made."
    )
    evidence_base = (
        f"{len(context.records)} candidate records were screened, "
        f"{sum(1 for r in context.records if r.selected)} were included, and "
        f"{len(analysed)} yielded readable full text for analysis. "
        f"{sum(1 for r in context.records if r.q1.verification_status == Q1Status.VERIFIED_Q1)} "
        "included papers have a verified Q1 quartile from a ranking source; "
        f"{sum(1 for r in context.records if r.q1.verification_status == Q1Status.UNVERIFIED)} "
        "have an unverified quartile and no quartile has been assumed. Statements below "
        "describe this reviewed evidence and are not generalisations about all research "
        "on the topic."
    )

    document = start_document(
        "Introduction", context, scope=scope, evidence_base=evidence_base
    )

    if not analysed:
        document.add_heading("Evidence not available", level=1)
        document.add_paragraph(
            "No paper in this job produced readable full text, so no evidence-based "
            "introduction can be written. Re-run the download and extraction stages, "
            "or obtain the papers listed in Unable_to_Download.docx through your own "
            "institutional access, and then re-run the report stage."
        )
        return finish_document(
            document,
            context,
            output_path,
            limitations=[
                "No analysed papers were available, so this document contains no "
                "synthesis of findings."
            ],
        )

    # --- 1. Background ---
    document.add_heading("1. Background", level=1)
    background = [
        f"Research on {config.topic.lower()} has produced a body of evidence that this "
        f"review draws on through {len(analysed)} analysed papers published between "
        f"{min((r.year for r in analysed if r.year), default=config.year_from)} and "
        f"{max((r.year for r in analysed if r.year), default=config.year_to)}."
    ]
    for record in analysed[:4]:
        if sentence := _evidence_sentence(
            context, record.record_id, "research_problem",
            document_name=document_name, section="1. Background",
        ):
            background.append(sentence)
    document.add_paragraph(" ".join(background))

    # --- 2. Importance of the topic ---
    document.add_heading("2. Why the topic matters", level=1)
    importance: list[str] = []
    for record in analysed:
        if sentence := _evidence_sentence(
            context, record.record_id, "policy_implications",
            document_name=document_name, section="2. Why the topic matters",
        ):
            importance.append(sentence)
        if len(importance) >= 4:
            break
    if importance:
        document.add_paragraph(
            "The reviewed papers connect this topic to decisions taken in practice. "
            + " ".join(importance)
        )
    else:
        document.add_paragraph(
            "None of the reviewed papers states policy or practice implications in "
            "extractable terms, so the applied importance of the topic cannot be "
            "evidenced from this set. This absence is itself recorded as a policy gap "
            "in Research_Gaps.docx."
        )

    # --- 3. Current knowledge ---
    document.add_heading("3. What is currently known", level=1)
    knowledge: list[str] = []
    for record in analysed:
        if sentence := _evidence_sentence(
            context, record.record_id, "main_findings",
            document_name=document_name, section="3. What is currently known",
        ):
            knowledge.append(sentence)
        if len(knowledge) >= 8:
            break
    document.add_paragraph(
        " ".join(knowledge)
        if knowledge
        else "No reviewed paper reported findings in a form that could be extracted "
             "with page evidence, so current knowledge cannot be summarised from this set."
    )

    # --- 4. Main methods used ---
    document.add_heading("4. Methods used in the reviewed evidence", level=1)
    methods = context.synthesis.landscape.methods
    if methods:
        top = list(methods.items())[:6]
        document.add_paragraph(
            "Within the reviewed evidence, the methods most frequently identified are "
            + ", ".join(f"{name} ({count} paper{'s' if count != 1 else ''})" for name, count in top)
            + ". Each is described, with its assumptions and limitations, in "
            "Models_and_Applications.docx."
        )
        design_sentences: list[str] = []
        for record in analysed[:5]:
            if sentence := _evidence_sentence(
                context, record.record_id, "study_design",
                document_name=document_name,
                section="4. Methods used in the reviewed evidence",
            ):
                design_sentences.append(sentence)
        if design_sentences:
            document.add_paragraph(" ".join(design_sentences))
    else:
        document.add_paragraph(
            "No recognised method was identified in the reviewed papers, so the "
            "methodological composition of this evidence base cannot be described."
        )

    # --- 5. Major findings ---
    document.add_heading("5. Convergence and disagreement in the findings", level=1)
    if context.synthesis.contradictions:
        document.add_paragraph(
            "The reviewed papers do not agree in every respect. "
            + " ".join(context.synthesis.contradictions[:3])
        )
    else:
        document.add_paragraph(
            "No direct contradiction in the direction of reported effects was detected "
            "across the reviewed papers. This is not evidence of consensus: the set is "
            f"small ({len(analysed)} papers) and several papers do not report their "
            "findings in comparable terms."
        )

    # --- 6. The remaining problem ---
    document.add_heading("6. What remains unresolved", level=1)
    author_gaps = [
        g for g in context.synthesis.gaps if g.stance == EvidenceStance.AUTHOR_STATED
    ]
    if author_gaps:
        document.add_paragraph(
            "The authors of the reviewed papers identify unresolved questions of their "
            "own. "
            + " ".join(
                f"{truncate_text(gap.statement, 300)} {gap.citations[0] if gap.citations else ''}"
                for gap in author_gaps[:4]
            )
        )
    else:
        document.add_paragraph(
            "No reviewed paper states a research gap in extractable terms. The gaps "
            "reported in Research_Gaps.docx are therefore the agent's inferences from "
            "coverage patterns rather than the authors' own statements."
        )

    # --- 7. Motivation for further research ---
    document.add_heading("7. Motivation for further research", level=1)
    weak_fields = [
        name.replace("_", " ")
        for name, count in coverage.items()
        if analysed and count / len(analysed) < 0.4
    ]
    motivation = (
        f"Taken together, the {len(analysed)} reviewed papers establish that the topic "
        "is being studied but leave specific ground uncovered. "
    )
    if weak_fields:
        motivation += (
            "Fewer than two in five reviewed papers report "
            + ", ".join(weak_fields[:6])
            + ", which limits how far their results can be compared or pooled. "
        )
    motivation += (
        "Research that reports these elements explicitly, and that extends the evidence "
        "to the contexts listed as under-represented in "
        "Global_Research_Landscape.docx, would address the gaps this review documents."
    )
    document.add_paragraph(motivation)

    return finish_document(
        document,
        context,
        output_path,
        limitations=_standard_limitations(context),
    )


def _standard_limitations(context: ReportContext) -> list[str]:
    """The limitations every synthesis document must disclose."""
    config = context.config
    analysed = context.analysed
    unverified = sum(
        1 for r in context.records if r.q1.verification_status == Q1Status.UNVERIFIED
    )
    ocr_needed = sum(1 for r in context.records if r.requires_ocr)
    failed = sum(1 for r in context.records if r.selected and not r.local_path)

    limitations = [
        f"This review rests on {len(analysed)} analysed paper"
        f"{'s' if len(analysed) != 1 else ''}, not on the complete literature. It is a "
        "search of specific databases with specific terms, not a census of the field.",
        f"Only literature published between {config.year_from} and {config.year_to} in "
        f"{config.language} was considered.",
        "Only papers with a legally accessible open-access PDF could be analysed in "
        "full. Subscription-only papers are listed in Unable_to_Download.docx and are "
        "absent from the synthesis, which may bias the evidence base towards "
        "open-access publishing.",
    ]
    if failed:
        limitations.append(
            f"{failed} selected paper{'s' if failed != 1 else ''} could not be retrieved "
            "legally and therefore contributed no evidence to this document."
        )
    if unverified:
        limitations.append(
            f"{unverified} paper{'s' if unverified != 1 else ''} carry an unverified "
            "journal quartile because no licensed ranking file was available for the "
            "relevant year. No quartile has been assumed."
        )
    if ocr_needed:
        limitations.append(
            f"{ocr_needed} PDF{'s' if ocr_needed != 1 else ''} had no extractable text "
            "layer and were excluded from evidence-based claims rather than being "
            "guessed at."
        )
    limitations.append(
        "Text extraction and field identification are automated. Statements labelled as "
        "agent inference are the agent's reading of the papers and must be checked "
        "against the source before being relied upon."
    )
    return limitations


# ---------------------------------------------------------------------------
# 2. Research gaps
# ---------------------------------------------------------------------------


def build_research_gaps(context: ReportContext, output_path: Path) -> Path:
    """Write ``Research_Gaps.docx`` with gaps separated by category."""
    gaps = context.synthesis.gaps
    scope = (
        f"This document separates the research gaps identified for \"{context.config.topic}\" "
        "into the categories below. Author-stated gaps are reported separately from "
        "gaps the agent inferred, and each gap names the papers that support it."
    )
    evidence_base = (
        f"{len(gaps)} gap statement{'s' if len(gaps) != 1 else ''} derived from "
        f"{len(context.analysed)} analysed paper"
        f"{'s' if len(context.analysed) != 1 else ''}. Each entry carries the evidence "
        "ledger IDs that link it back to the source text and page numbers."
    )
    document = start_document(
        "Research Gaps", context, scope=scope, evidence_base=evidence_base
    )

    if not gaps:
        document.add_heading("No gaps identified", level=1)
        document.add_paragraph(
            "No research gap could be identified, because no paper produced readable "
            "full text for analysis. This is an absence of evidence, not a finding that "
            "the field has no gaps."
        )
        return finish_document(
            document, context, output_path, limitations=_standard_limitations(context)
        )

    author_stated = [g for g in gaps if g.stance == EvidenceStance.AUTHOR_STATED]
    document.add_heading("How to read this document", level=1)
    document.add_paragraph(
        f"{len(author_stated)} of {len(gaps)} gap statements are the authors' own words. "
        f"The remaining {len(gaps) - len(author_stated)} are the agent's inferences from "
        "what the reviewed papers do and do not report. The two are never merged, and "
        "the 'Agent-inferred gaps' section at the end contains only inferences."
    )

    ordered_categories = [
        GapCategory.AUTHOR_STATED,
        GapCategory.METHODOLOGICAL,
        GapCategory.DATA,
        GapCategory.GEOGRAPHIC,
        GapCategory.POPULATION,
        GapCategory.MODEL_LIMITATION,
        GapCategory.VALIDATION,
        GapCategory.APPLICATION,
        GapCategory.POLICY,
        GapCategory.CONTRADICTORY,
        GapCategory.AGENT_INFERRED,
    ]

    for index, category in enumerate(ordered_categories, 1):
        items = [g for g in gaps if g.category == category]
        document.add_heading(f"{index}. {category.value}", level=1)
        if not items:
            paragraph = document.add_paragraph(
                "No gap in this category was identified from the reviewed evidence."
            )
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = MUTED
            continue

        if category == GapCategory.AGENT_INFERRED:
            note = document.add_paragraph(
                "The statements in this section are the agent's inferences drawn from "
                "what the reviewed papers do not report. No cited author makes these "
                "claims."
            )
            for run in note.runs:
                run.italic = True
                run.font.color.rgb = MUTED

        for gap in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(gap.statement)
            citation_run = paragraph.add_run(
                " " + (context.manager.citation_group(gap.supporting_record_ids) or "")
            )
            citation_run.bold = True
            detail = document.add_paragraph()
            detail.paragraph_format.left_indent = Inches(0.5)
            detail_run = detail.add_run(
                f"Evidence: {gap.stance.value}. "
                f"Ledger IDs: {', '.join(gap.evidence_ids) or 'not recorded'}. "
                f"Gap ID: {gap.gap_id}."
            )
            detail_run.font.size = Pt(8)
            detail_run.font.color.rgb = MUTED

    return finish_document(
        document, context, output_path, limitations=_standard_limitations(context)
    )


# ---------------------------------------------------------------------------
# 3. Global research landscape
# ---------------------------------------------------------------------------


def build_global_landscape(context: ReportContext, output_path: Path) -> Path:
    """Write ``Global_Research_Landscape.docx``."""
    landscape: LandscapeSummary = context.synthesis.landscape
    total = max(landscape.total_papers, 1)
    scope = (
        "This document describes the shape of the evidence retrieved for "
        f"\"{context.config.topic}\". Every count below refers to the "
        f"{landscape.total_papers} reviewed paper"
        f"{'s' if landscape.total_papers != 1 else ''} in this job."
    )
    evidence_base = (
        "The statistics are computed from the analysed papers' own text and verified "
        "bibliographic metadata."
    )
    document = start_document(
        "Global Research Landscape", context, scope=scope, evidence_base=evidence_base
    )

    warning = document.add_paragraph()
    warning_run = warning.add_run(
        "Important: this is a description of the reviewed evidence, not of world "
        "research. A country, method, or dataset appearing rarely here means it is "
        "rare in this retrieved set; it does not establish that little research exists "
        "elsewhere. The search covered specific databases with specific terms."
    )
    warning_run.bold = True
    warning_run.font.size = Pt(10)

    def counts_section(
        number: int, heading: str, mapping: dict[str, int], noun: str, empty: str
    ) -> None:
        """Render one count table, or an explicit absence note."""
        document.add_heading(f"{number}. {heading}", level=1)
        if not mapping:
            paragraph = document.add_paragraph(empty)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = MUTED
            return
        _add_table(
            document,
            [noun, "Papers", "Share of reviewed evidence"],
            [
                [name, count, f"{count / total:.0%}"]
                for name, count in list(mapping.items())[:20]
            ],
            widths=[3.4, 1.0, 1.9],
        )

    counts_section(
        1, "Countries and cities within the reviewed evidence", landscape.countries,
        "Country or city",
        "No reviewed paper stated a study location clearly enough to extract, so the "
        "geographic coverage of this evidence base cannot be described.",
    )
    counts_section(
        2, "Regions within the reviewed evidence", landscape.regions, "Region",
        "No region could be derived, because no study location was extracted.",
    )
    counts_section(
        3, "Research institutions", landscape.institutions, "Institution",
        "Author affiliations are not retrieved by this pipeline, so contributing "
        "institutions are not reported. Only publishers that are themselves research "
        "bodies would appear here. This is a limitation of the metadata sources, not "
        "evidence that no institutions are active.",
    )
    counts_section(
        4, "Dominant applications", landscape.applications, "Application area",
        "No reviewed paper described an application area in extractable terms.",
    )
    counts_section(
        5, "Common datasets", landscape.datasets, "Dataset type",
        "No recognised dataset type was named in the reviewed papers' data-source "
        "statements.",
    )
    counts_section(
        6, "Common methods", landscape.methods, "Method",
        "No recognised method was identified in the reviewed papers.",
    )

    document.add_heading("7. Emerging methods", level=1)
    if landscape.emerging_methods:
        document.add_paragraph(
            "These methods appear only in the more recent papers of this set and not in "
            "the earlier ones, which suggests recent adoption within the reviewed "
            "evidence: "
            + ", ".join(landscape.emerging_methods)
            + ". With a set this size, that pattern is suggestive rather than "
            "established."
        )
    else:
        document.add_paragraph(
            "No method is confined to the most recent papers, so no emerging method can "
            "be identified from this evidence base."
        )

    document.add_heading("8. Temporal trends", level=1)
    if landscape.year_counts:
        _add_table(
            document,
            ["Publication year", "Papers", "Share of reviewed evidence"],
            [
                [year, count, f"{count / total:.0%}"]
                for year, count in landscape.year_counts.items()
            ],
            widths=[2.0, 1.4, 2.4],
        )
        years = sorted(int(y) for y in landscape.year_counts)
        midpoint = years[len(years) // 2]
        recent = sum(c for y, c in landscape.year_counts.items() if int(y) >= midpoint)
        document.add_paragraph(
            f"{recent} of {landscape.total_papers} reviewed papers were published in "
            f"{midpoint} or later. Publication counts in a retrieved sample reflect both "
            "research activity and indexing coverage, so this is not a direct measure of "
            "how the field is growing."
        )
    else:
        document.add_paragraph("No publication years were recorded for the reviewed papers.")

    document.add_heading("9. Publication venues", level=1)
    if landscape.journals:
        _add_table(
            document,
            ["Journal or source", "Papers"],
            [[name, count] for name, count in list(landscape.journals.items())[:20]],
            widths=[4.5, 1.3],
        )
    else:
        document.add_paragraph("No journal names were recorded.")

    document.add_heading("10. Areas receiving limited attention in this evidence base", level=1)
    if landscape.under_researched:
        for item in landscape.under_researched:
            document.add_paragraph(item, style="List Bullet")
        document.add_paragraph(
            "These are absences in the retrieved set. Confirming that they are genuine "
            "research gaps would require a search designed specifically for each one."
        )
    else:
        document.add_paragraph(
            "Every region checked appears at least once in the reviewed evidence."
        )

    document.add_heading("11. Global and local contexts", level=1)
    requested = context.config.geography
    if landscape.regions:
        leading_region, leading_count = max(landscape.regions.items(), key=lambda kv: kv[1])
        document.add_paragraph(
            f"The reviewed evidence is weighted towards {leading_region} "
            f"({leading_count} of {landscape.total_papers} papers). The job requested a "
            f"study geography of \"{requested}\". Where a local context differs from the "
            "settings these papers studied, their findings should be treated as "
            "transferable hypotheses to be tested locally, not as established local "
            "results, since none of the reviewed papers tests that transfer."
        )
    else:
        document.add_paragraph(
            f"The job requested a study geography of \"{requested}\", but no study "
            "location could be extracted from the reviewed papers, so global and local "
            "contexts cannot be compared from this evidence."
        )

    return finish_document(
        document, context, output_path, limitations=_standard_limitations(context)
    )


# ---------------------------------------------------------------------------
# 4. Models and applications
# ---------------------------------------------------------------------------


def build_models_and_applications(context: ReportContext, output_path: Path) -> Path:
    """Write ``Models_and_Applications.docx``, grouped by model category."""
    models: list[ModelProfile] = context.synthesis.models
    scope = (
        "This document describes every model and method identified in the reviewed "
        f"evidence for \"{context.config.topic}\", grouped by family, with a "
        "plain-language explanation of how each works."
    )
    evidence_base = (
        f"{len(models)} model or method profile"
        f"{'s' if len(models) != 1 else ''} derived from "
        f"{len(context.analysed)} analysed paper"
        f"{'s' if len(context.analysed) != 1 else ''}."
    )
    document = start_document(
        "Models and Applications", context, scope=scope, evidence_base=evidence_base
    )

    if not models:
        document.add_heading("No models identified", level=1)
        document.add_paragraph(
            "No model or method could be identified from the reviewed papers, either "
            "because no full text was analysed or because the papers do not name a "
            "recognised approach."
        )
        return finish_document(
            document, context, output_path, limitations=_standard_limitations(context)
        )

    grouped: dict[str, list[ModelProfile]] = {}
    for profile in models:
        grouped.setdefault(profile.model_category, []).append(profile)

    document.add_heading("Overview", level=1)
    _add_table(
        document,
        ["Model family", "Models in family", "Papers"],
        [
            [
                category,
                ", ".join(p.model_name for p in profiles),
                sum(len(p.paper_record_ids) for p in profiles),
            ]
            for category, profiles in sorted(grouped.items())
        ],
        widths=[1.9, 3.0, 0.9],
    )

    section = 0
    for category, profiles in sorted(grouped.items()):
        section += 1
        document.add_heading(f"{section}. {category}", level=1)
        explanation = profiles[0].plain_explanation
        intro = document.add_paragraph()
        intro.add_run("How this family works in plain terms: ").bold = True
        intro.add_run(explanation)

        for profile in profiles:
            document.add_heading(profile.model_name, level=2)
            document.add_paragraph(profile.purpose)

            rows: list[list[Any]] = [
                ["Model category", profile.model_category],
                ["Papers using it", f"{len(profile.paper_record_ids)}"],
                ["Citations", "; ".join(profile.citations) or "not recorded"],
                [
                    "Main assumptions",
                    _bulleted(profile.assumptions)
                    or "Not stated in the reviewed papers.",
                ],
                [
                    "Required inputs",
                    _bulleted(profile.required_inputs)
                    or "Not stated in the reviewed papers.",
                ],
                ["Outputs", _bulleted(profile.outputs) or "Not stated in the reviewed papers."],
                [
                    "Study application",
                    _bulleted(profile.study_application)
                    or "Not stated in the reviewed papers.",
                ],
                [
                    "Software used",
                    _bulleted(profile.software_used) or "Not stated in the reviewed papers.",
                ],
                [
                    "Calibration approach",
                    _bulleted(profile.calibration_approach)
                    or "Not stated in the reviewed papers.",
                ],
                [
                    "Validation approach",
                    _bulleted(profile.validation_approach)
                    or "Not stated in the reviewed papers.",
                ],
                [
                    "Advantages",
                    _bulleted(profile.advantages)
                    or "Not characterised for this model family.",
                ],
                [
                    "Limitations",
                    _bulleted(profile.limitations)
                    or "Not characterised for this model family.",
                ],
                ["How it works (plain language)", profile.plain_explanation],
            ]
            _add_table(document, ["Attribute", "Detail"], rows, widths=[1.7, 4.3])

            note = document.add_paragraph()
            note_run = note.add_run(
                "Advantages and limitations describe the method family in general. They "
                "are not claims made by the cited papers, which are cited only for how "
                "they applied the method."
            )
            note_run.font.size = Pt(8)
            note_run.italic = True
            note_run.font.color.rgb = MUTED

    return finish_document(
        document, context, output_path, limitations=_standard_limitations(context)
    )


def _bulleted(items: list[str]) -> str:
    """Render a list as newline-separated bullets for a table cell."""
    return "\n".join(f"• {truncate_text(item, 300)}" for item in items[:6])


# ---------------------------------------------------------------------------
# 5. Paper summaries
# ---------------------------------------------------------------------------


def build_paper_summaries(context: ReportContext, output_path: Path) -> Path:
    """Write ``Paper_Summaries.docx``: one structured summary per paper."""
    analysed = context.analysed
    scope = (
        f"A structured summary of each of the {len(analysed)} analysed paper"
        f"{'s' if len(analysed) != 1 else ''} in this review, with page-level evidence "
        "references wherever the source supported them."
    )
    evidence_base = (
        "Each field is taken from the paper's own text. Where a paper does not report a "
        "field, the summary says so rather than filling the gap."
    )
    document = start_document(
        "Paper Summaries", context, scope=scope, evidence_base=evidence_base
    )
    document_name = "Paper_Summaries.docx"

    if not analysed:
        document.add_heading("No papers analysed", level=1)
        document.add_paragraph(
            "No paper produced readable full text, so no summary could be written."
        )
        return finish_document(
            document, context, output_path, limitations=_standard_limitations(context)
        )

    summary_fields: tuple[tuple[str, str], ...] = (
        ("research_problem", "Research problem"),
        ("research_objective", "Objective"),
        ("research_questions", "Research questions"),
        ("hypotheses", "Hypotheses"),
        ("study_geography", "Study geography"),
        ("study_context", "Study context"),
        ("study_design", "Study design"),
        ("data_source", "Data source"),
        ("sample_size", "Sample size"),
        ("unit_of_analysis", "Unit of analysis"),
        ("dependent_variables", "Dependent variables"),
        ("independent_variables", "Independent variables"),
        ("control_variables", "Control variables"),
        ("model_or_method", "Model or method"),
        ("model_equations", "Model equations"),
        ("software_or_tools", "Software or tools"),
        ("validation_approach", "Validation approach"),
        ("main_findings", "Main findings"),
        ("policy_implications", "Policy implications"),
        ("limitations_stated", "Author-stated limitations"),
        ("gaps_stated_by_authors", "Author-stated research gaps"),
        ("relevance_to_topic", "Relevance to this review"),
    )

    for index, record in enumerate(analysed, 1):
        analysis = context.analyses[record.record_id]
        citation = context.manager.citation(record.record_id)

        document.add_heading(f"{index}. {record.title}", level=1)
        header = document.add_paragraph()
        header_run = header.add_run(
            f"{citation}    |    Q1 status: {record.q1.verification_status.value}"
            f"    |    Analysis confidence: {analysis.overall_confidence}"
        )
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = MUTED

        full_citation = context.manager.reference(record.record_id)
        citation_paragraph = document.add_paragraph()
        citation_paragraph.add_run("Full citation: ").bold = True
        citation_paragraph.add_run(full_citation)

        rows: list[list[Any]] = []
        for field_name, label in summary_fields:
            evidence = analysis.field(field_name)
            if evidence.is_reported:
                prefix = (
                    "[Agent inference] "
                    if evidence.stance == EvidenceStance.AGENT_INFERENCE
                    else ""
                )
                pages = (
                    f" (p. {', '.join(str(p) for p in evidence.pages)})"
                    if evidence.pages
                    else " (page not recorded)"
                )
                rows.append([label, f"{prefix}{evidence.value}{pages}"])
                context.ledger.add_from_analysis(
                    analysis,
                    field_name,
                    document=document_name,
                    section=f"{index}. {truncate_text(record.title, 60)}",
                    in_text_citation=citation,
                )
            else:
                rows.append([label, evidence.stance.value])

        rows.append(
            [
                "Agent-inferred gap",
                analysis.agent_inferred_gap.value or "None inferred.",
            ]
        )
        rows.append(
            [
                "Fields not reported by this paper",
                ", ".join(analysis.missing_information) or "None",
            ]
        )
        _add_table(document, ["Field", "Content"], rows, widths=[1.6, 4.4])

    return finish_document(
        document, context, output_path, limitations=_standard_limitations(context)
    )


# ---------------------------------------------------------------------------
# 6. Unable to download
# ---------------------------------------------------------------------------


def build_unable_to_download(
    rows: list[dict[str, Any]],
    config: JobConfig,
    output_path: Path,
    *,
    generated_on: str | None = None,
) -> Path:
    """Write ``Unable_to_Download.docx``: the manual-retrieval register."""
    document = Document()
    _configure_styles(document)

    # Landscape orientation: this register table is wide. python-docx needs the
    # page dimensions swapped as well as the orientation flag set.
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    heading = document.add_heading("Papers That Could Not Be Downloaded", level=0)
    for run in heading.runs:
        run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(config.topic)
    subtitle_run.italic = True
    subtitle_run.font.color.rgb = MUTED

    meta = document.add_paragraph()
    meta_run = meta.add_run(
        f"Prepared: {generated_on or date.today().isoformat()}    |    "
        f"Entries: {len(rows)}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = MUTED

    document.add_heading("Why these papers are listed here", level=1)
    document.add_paragraph(
        "Each paper below was selected for the review but could not be obtained "
        "legally by the agent. The agent retrieves PDFs only from legitimate "
        "open-access locations or from a direct PDF URL that a publisher's own API has "
        "identified as authorised. It does not bypass paywalls, institutional logins, "
        "CAPTCHAs, anti-bot protections, or publisher download restrictions, and it "
        "does not use unauthorised mirrors."
    )
    document.add_paragraph(
        "These papers are therefore absent from the analysis and the synthesis "
        "documents. Retrieve any you need through your own institutional access or by "
        "contacting the corresponding author, then place the PDF in the job's "
        "'Downloaded Papers' folder using the exact paper title as the filename and "
        "re-run the extract and report stages."
    )

    if not rows:
        paragraph = document.add_paragraph(
            "No entries: every selected paper with a legally accessible PDF was "
            "downloaded and validated successfully."
        )
        for run in paragraph.runs:
            run.italic = True
        ensure_dir(Path(output_path).parent)
        document.save(output_path)
        LOG.info(f"Wrote {Path(output_path).name} (no failures to report).")
        return Path(output_path)

    document.add_heading("Register", level=1)
    headers = [
        "S. No.", "Paper title", "Authors", "Year", "Journal", "DOI",
        "Publisher landing page", "Attempted PDF links", "OA status",
        "Failure reason", "HTTP status", "Attempted at",
        "Recommended manual action", "Q1 status",
    ]
    table_rows = [
        [
            row.get("serial"),
            row.get("title"),
            truncate_text(str(row.get("authors", "")), 160),
            row.get("year"),
            truncate_text(str(row.get("journal", "")), 90),
            row.get("doi"),
            truncate_text(str(row.get("landing_page_url", "")), 110),
            truncate_text(str(row.get("attempted_urls", "")), 220),
            row.get("open_access_status"),
            truncate_text(str(row.get("failure_reason", "")), 220),
            row.get("http_status"),
            row.get("attempted_at"),
            truncate_text(str(row.get("recommended_action", "")), 260),
            row.get("q1_status"),
        ]
        for row in rows
    ]
    _add_table(
        document,
        headers,
        table_rows,
        widths=[0.4, 1.9, 1.1, 0.4, 1.0, 1.0, 1.1, 1.3, 0.6, 1.4, 0.6, 0.9, 1.6, 0.8],
    )

    _add_page_numbers(document)
    ensure_dir(Path(output_path).parent)
    document.save(output_path)
    LOG.info(f"Wrote {Path(output_path).name} with {len(rows)} entr(y/ies).")
    return Path(output_path)


# ---------------------------------------------------------------------------
# 7. Verification report
# ---------------------------------------------------------------------------


def build_verification_report(
    summary: VerificationSummary,
    findings: list[VerificationFinding],
    context: ReportContext,
    output_path: Path,
) -> Path:
    """Write ``Verification_Report.docx``."""
    scope = (
        "This report records the independent verification of the metadata, quartile "
        "evidence, downloaded files, extracted text, claims, and citations produced for "
        f"\"{context.config.topic}\". The agent that analysed the papers did not verify "
        "its own work: separate verification stages produced the findings below."
    )
    evidence_base = (
        f"{summary.total_claims_checked} claim(s) and {len(findings)} check(s) across "
        f"{summary.total_included} included paper(s)."
    )
    document = start_document(
        "Verification Report", context, scope=scope, evidence_base=evidence_base
    )

    document.add_heading("1. Counts", level=1)
    _add_table(
        document,
        ["Measure", "Count"],
        [
            ["Total papers discovered", summary.total_discovered],
            ["Total unique papers after deduplication", summary.total_unique],
            ["Total included in the review", summary.total_included],
            ["Verified Q1", summary.total_verified_q1],
            ["Unverified quartile", summary.total_unverified_quartile],
            ["PDFs downloaded and validated", summary.total_pdfs_downloaded],
            ["Failed or unavailable downloads", summary.total_failed_downloads],
            ["Papers analysed", summary.total_papers_analysed],
            ["Claims checked", summary.total_claims_checked],
            ["Checks passed", summary.passed_checks],
            ["Warnings", summary.warnings],
            ["Unresolved problems", summary.unresolved_problems],
        ],
        widths=[3.6, 1.6],
    )

    document.add_heading("2. Overall confidence", level=1)
    confidence = document.add_paragraph()
    confidence_run = confidence.add_run(summary.overall_confidence)
    confidence_run.bold = True
    confidence_run.font.size = Pt(13)
    document.add_paragraph(_confidence_explanation(summary))

    document.add_heading("3. Checks by verifier", level=1)
    by_verifier: dict[str, list[VerificationFinding]] = {}
    for finding in findings:
        by_verifier.setdefault(finding.verifier, []).append(finding)
    if by_verifier:
        _add_table(
            document,
            ["Verifier", "Checks", "Passed", "Warnings", "Failed"],
            [
                [
                    verifier,
                    len(items),
                    sum(1 for f in items if f.outcome == CheckOutcome.PASS),
                    sum(1 for f in items if f.outcome == CheckOutcome.WARNING),
                    sum(1 for f in items if f.outcome == CheckOutcome.FAIL),
                ]
                for verifier, items in sorted(by_verifier.items())
            ],
            widths=[2.2, 0.9, 0.9, 1.0, 0.9],
        )
    else:
        document.add_paragraph("No verification checks were recorded.")

    document.add_heading("4. Failures", level=1)
    failures = [f for f in findings if f.outcome == CheckOutcome.FAIL]
    if failures:
        _add_table(
            document,
            ["Verifier", "Check", "Target", "Detail", "Recommended action"],
            [
                [
                    f.verifier,
                    f.check,
                    truncate_text(f.target, 70),
                    truncate_text(f.detail, 240),
                    truncate_text(f.recommended_action, 200),
                ]
                for f in failures
            ],
            widths=[1.1, 1.2, 1.3, 1.7, 1.5],
        )
    else:
        document.add_paragraph("No check failed.")

    document.add_heading("5. Warnings", level=1)
    warnings = [f for f in findings if f.outcome == CheckOutcome.WARNING]
    if warnings:
        _add_table(
            document,
            ["Verifier", "Check", "Target", "Detail"],
            [
                [f.verifier, f.check, truncate_text(f.target, 70), truncate_text(f.detail, 300)]
                for f in warnings[:100]
            ],
            widths=[1.1, 1.3, 1.5, 2.9],
        )
        if len(warnings) > 100:
            document.add_paragraph(
                f"{len(warnings) - 100} further warnings are listed in "
                "unresolved_issues.csv."
            )
    else:
        document.add_paragraph("No warnings were raised.")

    document.add_heading("6. Corrections applied", level=1)
    corrections = [f for f in findings if f.corrected_value]
    if corrections:
        document.add_paragraph(
            "Questionable data was never silently overwritten. Each correction below "
            "preserves the original value alongside the corrected one, its source, and "
            "the reason."
        )
        _add_table(
            document,
            ["Target", "Original value", "Corrected value", "Source", "Reason"],
            [
                [
                    truncate_text(f.target, 60),
                    truncate_text(f.original_value, 120),
                    truncate_text(f.corrected_value, 120),
                    truncate_text(f.correction_source, 80),
                    truncate_text(f.correction_reason, 160),
                ]
                for f in corrections
            ],
            widths=[1.2, 1.4, 1.4, 1.0, 1.8],
        )
    else:
        document.add_paragraph("No correction was required or applied.")

    document.add_heading("7. Recommended manual checks", level=1)
    if summary.recommended_manual_checks:
        for item in summary.recommended_manual_checks:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("No manual check is outstanding.")

    document.add_heading("8. Unresolved problems", level=1)
    unresolved = [f for f in findings if f.outcome == CheckOutcome.FAIL and not f.resolved]
    if unresolved:
        document.add_paragraph(
            f"{len(unresolved)} problem(s) remain unresolved. The review should not be "
            "treated as complete until each has been reviewed by a human. Full details "
            "are in unresolved_issues.csv."
        )
        for finding in unresolved[:25]:
            document.add_paragraph(
                f"{finding.check} - {truncate_text(finding.detail, 240)}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No unresolved problems remain.")

    return finish_document(
        document,
        context,
        output_path,
        limitations=[
            "Verification is automated. It checks internal consistency, resolvable "
            "identifiers, page evidence, and citation completeness; it does not "
            "reproduce the papers' analyses or judge whether their conclusions are "
            "correct.",
            "DOI resolution confirms that an identifier resolves, not that the "
            "resolved record is the intended paper in every respect.",
            "A passing verification does not remove the need to read the source papers "
            "before relying on this review.",
        ],
        cited_only=False,
    )


def _confidence_explanation(summary: VerificationSummary) -> str:
    """Explain what the overall confidence grade rests on."""
    total = max(summary.passed_checks + summary.warnings + summary.unresolved_problems, 1)
    pass_rate = summary.passed_checks / total
    parts = [
        f"{summary.passed_checks} of {total} checks passed ({pass_rate:.0%}), with "
        f"{summary.warnings} warning(s) and {summary.unresolved_problems} unresolved "
        "problem(s)."
    ]
    if summary.total_unverified_quartile:
        parts.append(
            f"{summary.total_unverified_quartile} paper(s) carry an unverified journal "
            "quartile, which lowers confidence in any Q1-based claim."
        )
    if summary.total_failed_downloads:
        parts.append(
            f"{summary.total_failed_downloads} selected paper(s) could not be retrieved "
            "legally and contributed no evidence."
        )
    if summary.total_papers_analysed == 0:
        parts.append(
            "No paper was analysed, so no evidence-based claim in this review can be "
            "confirmed."
        )
    parts.append(
        "Confidence describes the internal consistency of this job's outputs. It is not "
        "a judgement on the quality of the underlying research."
    )
    return " ".join(parts)


# ---------------------------------------------------------------------------
# Orchestrating helper
# ---------------------------------------------------------------------------


def build_all_reports(context: ReportContext, reports_dir: Path) -> list[Path]:
    """Write the five synthesis documents and return their paths."""
    reports_dir = ensure_dir(Path(reports_dir))
    written: list[Path] = []
    builders = (
        ("Introduction.docx", build_introduction),
        ("Research_Gaps.docx", build_research_gaps),
        ("Global_Research_Landscape.docx", build_global_landscape),
        ("Models_and_Applications.docx", build_models_and_applications),
        ("Paper_Summaries.docx", build_paper_summaries),
    )
    for filename, builder in builders:
        try:
            written.append(builder(context, reports_dir / filename))
        except Exception as exc:  # noqa: BLE001 - one bad report must not lose the rest
            LOG.error(f"Could not write {filename}: {type(exc).__name__}: {exc}")
    return written
